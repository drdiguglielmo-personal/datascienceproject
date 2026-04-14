#!/usr/bin/env python3
"""Fill out the Part 2 report template with project results."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

PROJECT_DIR = Path(__file__).resolve().parents[1]
TEMPLATE = PROJECT_DIR / "Part2_Report" / "CSE40467 - Project Part 2 Report Template.docx"
OUTPUT = PROJECT_DIR / "Part2_Report" / "CSE40467 - Project Part 2 Report.docx"
FIGURES = PROJECT_DIR / "figures"


def set_run_font(run, name="Times New Roman", size=Pt(11), bold=False):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is not None:
        rFonts.set(qn("w:eastAsia"), name)


def add_paragraph_to_cell(cell, text, first=False):
    """Add one formatted paragraph to a cell. Supports **bold** markers."""
    if first:
        p = cell.paragraphs[0]
        p.clear()
    else:
        p = cell.add_paragraph()

    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if is_bold else part
        run = p.add_run(content)
        set_run_font(run, bold=is_bold)

    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(4)
    return p


def add_figure_to_cell(cell, fig_name, width_inches=4.5):
    """Add a placeholder line for a figure instead of embedding the image."""
    p = cell.add_paragraph()
    run = p.add_run(f"[Insert Figure: figures/{fig_name}]")
    set_run_font(run, bold=False)
    run.font.italic = True
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(4)


def fill_cell(cell, paragraphs, figures=None):
    """Fill a cell: clear old content, add paragraphs and optional figures."""
    while len(cell.paragraphs) > 1:
        el = cell.paragraphs[-1]._element
        el.getparent().remove(el)
    for i, text in enumerate(paragraphs):
        add_paragraph_to_cell(cell, text, first=(i == 0))
    if figures:
        for fig_name, w in figures:
            add_figure_to_cell(cell, fig_name, w)


# ---------------------------------------------------------------------------
# Content for each box
# ---------------------------------------------------------------------------

EXEC_SUMMARY = [
    (
        "In this project, we investigate the prediction of FIFA Men's World Cup match "
        "outcomes as a three-class classification problem: home team win, away team win, "
        "or draw. Using 964 matches spanning 22 World Cup tournaments from 1930 to 2022, "
        "we engineered 39 features that capture team strength, recent form, head-to-head "
        "records, and contextual information, all computed in a leakage-safe manner using "
        "only pre-match data. We evaluated seven supervised learning algorithms (K-Nearest "
        "Neighbors, Decision Tree, Naive Bayes, Support Vector Machine, Random Forest, "
        "Neural Network, and XGBoost) using temporal walk-forward cross-validation, which "
        "we found gives significantly more honest performance estimates than standard "
        "stratified cross-validation. Our best model, a tuned Random Forest with 300 trees "
        "and regularization through minimum leaf size, achieves 62.5% test accuracy and "
        "0.578 macro F1 on the held-out 2022 World Cup, exceeding the 50.0% majority-class "
        "baseline by 12.5 percentage points."
    ),
    (
        "Beyond the main model comparison, we conducted an extensive experimental phase "
        "in which we tested five external data sources (international match results, FIFA "
        "world rankings, qualifying records, squad market values, and StatsBomb expected "
        'goals). We found that on our 900-row training set, hyperparameter tuning through '
        "regularization delivered roughly four times more improvement than the best feature "
        'addition. We also discovered what we call the "overconfidence" pattern, in which '
        "more precise team-strength features paradoxically destroy draw prediction by making "
        "the model too confident about which team will win."
    ),
]

PROBLEM_DEF = [
    (
        "Predicting football match outcomes is a well-studied yet persistently difficult "
        "problem in sports analytics. The FIFA Men's World Cup is the most widely watched "
        "sporting event globally, and accurate predictions are of interest to researchers, "
        "sports analysts, and the general public. We chose this problem for several reasons."
    ),
    (
        "First, football is inherently a low-scoring, high-variance sport where upsets "
        "happen frequently. Unlike basketball or baseball, where large scoring margins "
        "smooth out randomness, a single goal can decide a football match regardless of "
        "which team dominated possession or created more chances. This makes prediction "
        "challenging and imposes a natural ceiling on any model's accuracy."
    ),
    (
        "Second, the three-class nature of the target (home win, away win, or draw) adds "
        "significant complexity compared to binary classification. Draws are particularly "
        "hard to predict because they emerge from many different situations: tactical "
        "decisions in group-stage matches where both teams need a point, evenly matched "
        "opponents, or fatigue late in a tournament. Published research by Hvattum and "
        "Arntzen (2010) reports roughly 53% to 55% accuracy for three-class prediction on "
        "international football, which means that even modest improvements above that "
        "range represent meaningful progress."
    ),
    (
        "Third, the dataset is small by modern machine learning standards. With only 900 "
        "training matches from 22 World Cups held over nearly a century, we face real "
        "constraints on model complexity and the number of features we can use before "
        "overfitting becomes a concern. This makes the project an instructive case study "
        "in how to apply data science techniques when data is scarce and every modeling "
        "decision matters."
    ),
    (
        "We believe this problem is worth studying because it highlights fundamental "
        "tensions in applied machine learning: the tradeoff between model complexity and "
        "data size, the importance of domain-appropriate evaluation (temporal "
        "cross-validation versus random splitting), and the surprising finding that more "
        "data and more features do not always improve predictions."
    ),
]

RELATED_WORK = [
    (
        "Predicting football match outcomes has been an active research area for several "
        "decades. Maher (1982) introduced one of the earliest statistical models, treating "
        "goal counts as independent Poisson processes with team-specific attack and defense "
        "parameters. His work established the idea that team strength can be decomposed "
        "into offensive and defensive components, a principle that influenced our feature "
        "design: we include separate goals-scored and goals-conceded features for each "
        "team. Dixon and Coles (1997) extended the Poisson model by adding a correlation "
        "correction for low-scoring outcomes, which improved draw prediction. Their "
        "observation that draws are the hardest outcome to model remains a central "
        "challenge in the field, and one we encountered directly in our own experiments."
    ),
    (
        "Hvattum and Arntzen (2010) conducted a systematic evaluation of ELO ratings as "
        "predictors for international football. Using over 10,000 matches, they showed "
        "that ELO-based ordinal logistic regression achieved roughly 53% to 55% accuracy "
        "for three-class prediction. Their work established ELO difference as the single "
        "strongest pre-match predictor, a finding that our project independently confirms: "
        "in our feature importance analysis, elo_diff consistently ranks as the most "
        "important feature."
    ),
    (
        "Hubacek, Sourek, and Zelezny (2019) developed a score-based prediction framework "
        "that exploited the ordinal structure of football outcomes (the ordering from away "
        "win through draw to home win). Their competitive results in the 2018 Soccer "
        "Prediction Challenge suggest that ordinal approaches may offer advantages over "
        "standard multiclass classifiers, a direction we note for future work."
    ),
    (
        "Groll, Ley, Schauberger, and Lock (2019) organized a formal prediction "
        "tournament for the 2018 World Cup with 26 participating research teams. The best "
        "methods combined ensemble learning, hybrid models, and diverse features. A key "
        "finding was that no single method dominates; instead, thoughtful feature "
        "engineering and honest evaluation matter more than model complexity. Our results "
        "align with this: our tuned Random Forest, a relatively straightforward ensemble, "
        "outperforms the more complex XGBoost. Additionally, the importance of proper "
        "temporal evaluation that Groll et al. highlighted is confirmed by our finding that "
        "standard stratified cross-validation overestimates performance by 4 to 9 "
        "percentage points compared to temporal walk-forward cross-validation."
    ),
    (
        "In the broader literature on football prediction with machine learning, ensemble "
        "methods such as random forests and gradient boosting have been found to "
        "consistently outperform simpler models like logistic regression, Naive Bayes, "
        "and k-nearest neighbors. Our results are consistent with this pattern: Random "
        "Forest and XGBoost were our top two models, while Naive Bayes performed worst. "
        "The literature also suggests that combining multiple feature sources tends to "
        "improve accuracy; however, our experiments reveal a nuance: on small datasets, "
        "additional features can hurt performance through overfitting, a finding with "
        "practical implications for how researchers approach feature engineering in "
        "data-scarce domains."
    ),
]

DATA_COLLECTION = [
    (
        "Our primary dataset comes from the Fjelstul World Cup Database v1.2.0 (Fjelstul, "
        "2021), a comprehensive open-source database of all FIFA World Cup matches. We used "
        "two raw CSV files from this source: matches.csv, which contains all World Cup "
        "matches for both men's and women's tournaments, and tournaments.csv, which provides "
        "metadata including the year of each tournament. These files were downloaded from "
        "the Fjelstul GitHub repository (https://github.com/jfjelstul/worldcup)."
    ),
    (
        "After filtering to men's World Cup matches only, the dataset contains 964 matches "
        "across 22 tournaments spanning from 1930 to 2022, featuring 84 distinct national "
        "teams. Each match record includes 38 columns covering tournament and match "
        "identifiers, stage information (group stage or knockout stage), team names and "
        "codes, match scores, extra-time and penalty indicators, venue details (stadium "
        "name, city, and country), and outcome labels."
    ),
    (
        "For our experimental phase, we additionally integrated five external data sources "
        "to test whether they could improve predictions: (1) 49,287 international football "
        "match results from 1872 to 2026, sourced from a Kaggle dataset by Mart Jurisoo; "
        "(2) 67,894 FIFA world ranking records from 1992 to 2024, obtained from Dato-Futbol "
        "on GitHub; (3) Transfermarkt player data consisting of 47,702 player records and "
        "616,377 historical market valuations from 2004 to 2026; (4) 13,843 World Cup squad "
        "roster entries from the Fjelstul repository; and (5) 256 team-match event records "
        "from StatsBomb open data covering the 2018 and 2022 World Cups."
    ),
    (
        "All data sources are publicly available under open licenses (CC0 Public Domain for "
        "the Kaggle datasets, and open access for the others). No web scraping or API "
        "access was required; all datasets were downloaded as CSV files."
    ),
]

DATA_PREPROCESSING = [
    (
        "Our preprocessing pipeline has two main stages: data cleaning and feature "
        "engineering."
    ),
    (
        "**Data Cleaning (clean_worldcup.py, 173 lines).** First, we filter the raw "
        "matches file to keep only men's World Cup matches by matching on the tournament "
        "name field, and then join with the tournament table to obtain the year column. "
        "Second, we perform type cleaning: date columns are parsed using pandas "
        "to_datetime with errors='coerce' to safely handle any malformed values; nine "
        "boolean-like columns (indicating group stage, knockout stage, extra time, "
        "penalties, etc.) are cast from object type to integer using pd.to_numeric with "
        "fillna(0) for defensive handling; and six score columns are ensured to be "
        "numeric. Third, we split the data by tournament year: matches from 1930 to 2018 "
        "form the training set (900 matches), and matches from 2022 form the test set (64 "
        "matches). This split simulates the real-world scenario where a model trained on "
        "all historical data is used to predict the next World Cup."
    ),
    (
        "**Feature Engineering (feature_engineering.py, 1,985 lines).** This script "
        "concatenates the training and test sets, sorts them chronologically, and computes "
        "39 features. A critical design decision is the use of a date-level shift rather "
        "than a row-level shift when computing expanding-window features. Because multiple "
        "World Cup matches can occur on the same day, a row-level shift could allow "
        "same-day results to leak into features. Our date-level shift prevents this. "
        "After computing all features, missing values are filled with principled defaults "
        "(described in the Variables section), and the data is split back into separate "
        "training and test CSV files. Post-match columns such as scores and outcomes are "
        "explicitly excluded from the output to prevent accidental leakage."
    ),
    (
        "**Missing Value Strategy.** Rate features (win rate, draw rate) receive a fill "
        "value of 0.33, representing a uniform prior for the three-class problem. Count "
        "features (matches played, appearances) receive 0 to indicate no prior history. "
        "ELO ratings start at 1500 (the standard initial value). Rest days are filled "
        "with the median of existing values. Interaction features are filled with 0.0, "
        "consistent with the zero-filled components they are derived from."
    ),
]

DATA_DOCUMENTATION = [
    (
        "The cleaned dataset consists of 964 matches with 38 raw columns per match. The "
        "training set contains 900 matches from 21 tournaments (1930 to 2018), and the test "
        "set contains 64 matches from the 2022 Qatar World Cup. There are 84 distinct "
        "national teams across the full dataset. A detailed codebook documenting every "
        "column is maintained in docs/worldcup_subset_codebook.csv."
    ),
    (
        "**Key descriptive statistics.** Home teams score an average of 1.78 goals per "
        "match compared to 1.05 for away teams. There are no null values in the cleaned "
        "dataset. The raw columns include: tournament_id and match_id (identifiers), "
        "group_stage and knockout_stage (binary indicators), stage_name (text label for "
        "the round), home_team_name and away_team_name, team codes, home_team_score and "
        "away_team_score, extra_time and penalty indicators, stadium_name, city_name, "
        "country_name, and the outcome labels (result, home_team_win, away_team_win, draw)."
    ),
    (
        "**Target variable distribution.** The target is \"result\" with three classes. In "
        "the training set: home team win appears 513 times (57.0%), away team win appears "
        "218 times (24.2%), and draw appears 169 times (18.8%). In the test set, the "
        "distribution is more balanced: 32 home wins (50.0%), 16 away wins (25.0%), and "
        "16 draws (25.0%). This distributional shift between training and test is notable "
        "because models that heavily exploit the home-win prior will be penalized on the "
        "more balanced 2022 test set."
    ),
    (
        "**Important domain note.** The \"home team\" designation in World Cup matches is a "
        "FIFA administrative assignment at a neutral venue, not a true home advantage as "
        "in domestic league matches. This is a key nuance of the dataset that our "
        "home_is_host and continent advantage features attempt to address."
    ),
    (
        "After feature engineering, the final processed dataset has 39 features plus "
        "metadata and target columns per match."
    ),
]

VARIABLES = [
    (
        '**Target variable (dependent variable).** The target is "result," a three-class '
        "categorical variable indicating the match outcome: home team win, away team win, "
        "or draw."
    ),
    (
        "**Feature variables (independent variables).** Our model uses 39 engineered "
        "features organized into 10 categories, all computed using only information "
        "available before each match:"
    ),
    (
        "(A) **Team Historical Performance** (12 features): cumulative win rate, draw "
        "rate, goals per game, goals conceded per game, and total matches played for both "
        "home and away teams, plus difference features for win rate and goals. These use "
        "expanding-window aggregates over all prior World Cup matches."
    ),
    (
        "(B) **Match Context** (3 features): binary indicators for group stage and "
        "knockout stage, plus an ordinal encoding of the tournament stage (group = 0 "
        "through final = 6)."
    ),
    (
        "(C) **Host Advantage** (2 features): binary flags for whether the home or away "
        "team is the tournament host nation."
    ),
    (
        "(D) **World Cup Experience** (3 features): count of distinct prior World Cup "
        "tournaments each team appeared in, plus their difference."
    ),
    (
        "(E) **Head-to-Head Record** (5 features): total prior World Cup meetings between "
        "the two teams, wins for each side, draws, and the home team's win rate in those "
        "meetings (defaulting to 0.33 for first-time matchups)."
    ),
    (
        "(F) **ELO Ratings** (3 features): pre-match ELO for each team (all teams start "
        "at 1500, K-factor of 32) and their difference. Computed from World Cup matches "
        "only."
    ),
    (
        "(G) **Rolling Form** (4 features): win rate and goals per game over each team's "
        "last 5 World Cup matches, with fallbacks for teams with fewer prior matches."
    ),
    "(H) **Rest Days** (2 features): days since each team's previous World Cup match, capped at 365.",
    (
        "(I) **Interaction Features** (3 features): home attack multiplied by away "
        "defense, away attack multiplied by home defense, and ELO difference multiplied "
        "by rolling form difference. These capture non-linear relationships."
    ),
    (
        "(J) **Home Continent Advantage** (2 features): binary flags indicating whether "
        "each team's FIFA confederation matches the host country's confederation. This "
        "captures the documented effect that European teams win over 60% of World Cup "
        "matches played in Europe, and similarly for South American teams at South "
        "American World Cups."
    ),
]

ALGO_1_RF = [
    "**Algorithm 1: Random Forest (Best Model)**",
    (
        "Our best-performing model is a Random Forest classifier, an ensemble method that "
        "aggregates predictions from many decision trees, each trained on a bootstrap "
        "sample of the data. We chose Random Forest as our primary model because "
        "bagging-based ensembles are effective at reducing variance, which is particularly "
        "important on our small 900-row dataset where individual trees tend to overfit."
    ),
    (
        "**Final hyperparameters:** n_estimators = 300 (number of trees), max_features "
        "= 'sqrt' (number of features considered at each split), min_samples_leaf = 5 "
        "(minimum samples required in a terminal node), and class_weight = 'balanced' "
        "(class weights inversely proportional to frequency)."
    ),
    (
        "**Hyperparameter tuning process.** We determined these parameters through a "
        "grid search using 4-fold temporal walk-forward cross-validation. The search "
        "explored n_estimators in {100, 200, 300, 500}, max_features in {0.2, 0.3, "
        "'sqrt', 'log2'}, and min_samples_leaf in {1, 3, 5}. The most impactful "
        "parameter was min_samples_leaf: increasing it from the default of 1 to 5 "
        "improved cross-validation F1 from 0.488 to 0.531, a gain of 4.3 percentage "
        "points. This was the single largest improvement in the entire project, exceeding "
        "the benefit of any feature addition. It works by preventing the model from "
        "creating terminal nodes that memorize noise in small subsets of the training "
        "data, effectively acting as regularization."
    ),
    (
        "**Class imbalance handling.** The class_weight = 'balanced' setting addresses "
        "the imbalance in our training data (57% home wins, 24% away wins, 19% draws) "
        "by assigning higher weight to underrepresented classes. Without this, the model "
        "tends to ignore the minority draw class almost entirely."
    ),
    (
        "**Results:** test accuracy of 62.5%, test macro F1 of 0.578, test RPS of 0.133, "
        "and temporal cross-validation F1 of 0.541. The top features by importance are "
        "elo_diff, home_elo, away_elo, elo_x_form_diff, hist_win_rate_diff, and "
        "home_hist_win_rate."
    ),
]

ALGO_2_XGB = [
    "**Algorithm 2: XGBoost (Extreme Gradient Boosting)**",
    (
        "XGBoost is our second model, chosen because it represents a fundamentally "
        "different ensemble strategy from Random Forest. While Random Forest uses bagging "
        "(training independent trees on bootstrap samples and averaging them), XGBoost "
        "uses boosting: it builds trees sequentially, with each new tree correcting the "
        "errors of the previous ones through gradient descent on a loss function. We "
        "wanted to test whether this sequential error-correction approach would outperform "
        "the parallel averaging of Random Forest on our dataset."
    ),
    (
        "**Hyperparameters.** We tested three configurations via temporal "
        "cross-validation, varying max_depth (3, 4, 5), n_estimators (100, 150, 200), "
        "and learning_rate (0.05, 0.08, 0.1). We applied L1 and L2 regularization "
        "(reg_alpha = 0.5, reg_lambda = 2.0) to control overfitting and gave draws 1.5 "
        "times sample weight to address the class imbalance."
    ),
    (
        "**Results and comparison with Random Forest.** XGBoost achieved a test accuracy "
        "of 59.4% and macro F1 of 0.514, compared to Random Forest's 62.5% and 0.578. "
        "The temporal cross-validation F1 was 0.434 versus Random Forest's 0.541. We "
        "believe this underperformance is primarily a dataset size effect. Boosting "
        "methods generally need larger training sets to realize their advantages because "
        "each sequential tree has to learn meaningful residual patterns from the previous "
        "trees' errors. With only 900 training rows, the residual patterns tend to be "
        "dominated by noise rather than genuine signal. Random Forest's variance "
        "reduction through averaging many independent trees proves more effective at this "
        "scale. This finding is consistent with the general observation in the machine "
        "learning literature that bagging tends to outperform boosting on smaller "
        "datasets."
    ),
]

ALGO_3_OTHERS = [
    "**Algorithm 3: Five Additional Models**",
    (
        "We evaluated five more algorithms to provide a comprehensive comparison across "
        "different modeling paradigms."
    ),
    (
        "**K-Nearest Neighbors (KNN).** KNN is a non-parametric method that classifies "
        "each match by majority vote among its k nearest neighbors in feature space. We "
        "tuned k over {3, 5, 7, 9} using temporal cross-validation. As a distance-based "
        "method, it benefits from the StandardScaler normalization we applied. KNN serves "
        "as a simple baseline that makes no assumptions about the data distribution. It "
        "achieved 53.1% test accuracy and 0.463 macro F1."
    ),
    (
        "**Decision Tree.** A single decision tree with max_depth tuned over "
        "{3, 5, 7, 10, None}. We included it primarily for its interpretability and the "
        "feature importance rankings it produces, which helped us understand what "
        "information drives predictions. Its analysis confirmed elo_diff as the top "
        "predictor. Test accuracy: 51.6%, test F1: 0.469."
    ),
    (
        "**Gaussian Naive Bayes.** This probabilistic method assumes that features are "
        "conditionally independent given the class label. We included it knowing this "
        "assumption is violated by our correlated features (for example, home_elo and "
        "home_hist_win_rate both measure team strength). It performed worst of all models "
        "with 32.8% test accuracy and 0.316 macro F1, even below the majority-class "
        "baseline, demonstrating the cost of the independence assumption when features "
        "carry overlapping information."
    ),
    (
        "**SVM with RBF Kernel.** We tuned the regularization parameter C over "
        "{0.1, 1, 10, 100} using temporal cross-validation and set class_weight = "
        "'balanced.' SVM achieved 51.6% test accuracy and 0.392 F1. Notably, it had "
        "the second-best RPS (0.139), suggesting that its predicted probabilities are "
        "reasonably well-calibrated even when its point predictions are not the most "
        "accurate."
    ),
    (
        "**Neural Network (MLP).** A multi-layer perceptron with two hidden layers of "
        "64 and 32 neurons, early stopping, and a maximum of 500 iterations. It achieved "
        "51.6% test accuracy and 0.415 F1. The MLP showed considerable instability: its "
        "temporal cross-validation accuracy was only 41.1%, much lower than its test "
        "accuracy of 51.6%. This variability suggests that 900 training samples is not "
        "enough for neural networks to learn stable patterns, even with early stopping "
        "as regularization."
    ),
]

BASELINE = [
    (
        'Our baseline is the majority-class classifier, which always predicts "home team '
        'win" regardless of any match information. Since home team wins account for 57.0% '
        "of the training data, this baseline achieves 57.0% accuracy on the training set. "
        "On the 2022 test set, where the distribution is more balanced (50.0% home wins, "
        "25.0% away wins, 25.0% draws), the baseline achieves 50.0% accuracy."
    ),
    (
        "We chose this baseline because it represents the simplest possible prediction "
        "strategy and sets a meaningful bar that any useful model must exceed. It is more "
        "demanding than a random classifier (which would achieve roughly 33.3% on a "
        "balanced three-class problem) and better reflects what a naive human predictor "
        'might do: simply guess that the stronger team (as indicated by FIFA\'s "home" '
        "label) will win."
    ),
    (
        "Our best model, the tuned Random Forest, surpasses this baseline by 12.5 "
        "percentage points on the test set (62.5% versus 50.0%), demonstrating that the "
        "features we engineered carry genuine predictive signal beyond what a naive "
        "majority prediction offers."
    ),
]

METRICS = [
    (
        "We used three primary metrics to evaluate all models, along with confusion "
        "matrices and per-class classification reports."
    ),
    (
        "**Accuracy** measures the fraction of correct predictions. While intuitive, it "
        "can be misleading on imbalanced datasets, since a model that ignores minority "
        "classes can still appear accurate."
    ),
    (
        "**Macro F1** is the unweighted average of per-class F1 scores. It penalizes "
        "models that perform poorly on any class, making it our most informative single "
        "metric for this imbalanced three-class problem. A model that achieves high "
        "accuracy by ignoring draws would receive a low macro F1."
    ),
    (
        "**Ranked Probability Score (RPS)** measures the quality of predicted probability "
        "distributions assuming an ordinal ordering (away win < draw < home win). Lower "
        "is better. This metric is standard in sports forecasting and evaluates whether "
        "a model's confidence is well-calibrated, not just whether its top prediction is "
        "correct."
    ),
    "**Results on the 2022 World Cup test set (64 matches):**",
    "Random Forest (tuned): 62.5% accuracy, 0.578 F1, 0.133 RPS (best on all three metrics).",
    "XGBoost: 59.4% accuracy, 0.514 F1, 0.146 RPS.",
    "KNN: 53.1% accuracy, 0.463 F1, 0.158 RPS.",
    "Decision Tree: 51.6% accuracy, 0.469 F1, 0.165 RPS.",
    "SVM (RBF): 51.6% accuracy, 0.392 F1, 0.139 RPS.",
    "Neural Network: 51.6% accuracy, 0.415 F1, 0.148 RPS.",
    "Naive Bayes: 32.8% accuracy, 0.316 F1, 0.253 RPS.",
    (
        "Under temporal cross-validation, every model showed 4 to 9 percentage points "
        "lower accuracy compared to standard stratified cross-validation. This confirms "
        "that shuffled cross-validation overestimates performance on temporal sports data "
        "by allowing the model to train on future tournaments and validate on past ones."
    ),
    (
        "The best model's draw recall was approximately 44% in cross-validation, "
        "consistent with published findings that draws are the hardest outcome to "
        "predict. The top features by Random Forest importance were elo_diff, home_elo, "
        "away_elo, and elo_x_form_diff."
    ),
]

DISCUSSION = [
    (
        "Our tuned Random Forest with 39 features achieves 62.5% test accuracy and 0.578 "
        "macro F1 on the 2022 World Cup. The temporal cross-validation F1 of 0.541, "
        "computed across roughly 256 validation matches from four tournament-year folds, "
        "gives a more stable performance estimate. Both numbers are competitive with "
        "published benchmarks of 53% to 55% accuracy for three-class international "
        "football prediction (Hvattum and Arntzen, 2010). We note that our 64-match test "
        "set carries high variance: a single upset can shift accuracy by about 1.5 "
        "percentage points."
    ),
    (
        "**The value of regularization over feature engineering.** The most impactful "
        "improvement in our project came not from adding features or data, but from "
        "regularization. Increasing min_samples_leaf from 1 to 5 improved "
        "cross-validation F1 by 4.3 percentage points, roughly four times more than the "
        "best feature addition (continent advantage at 1.0 percentage point). This tells "
        "us that on a 900-row dataset, the model was overfitting to noise rather than "
        "lacking information. This result connects to the bias-variance tradeoff: on "
        "small datasets, reducing model complexity tends to be more effective than adding "
        "new signals."
    ),
    (
        '**The "overconfidence" pattern.** We discovered that features which more '
        "precisely measure team strength (such as international ELO, FIFA ranking points, "
        "and squad market values) paradoxically hurt the model. These features make the "
        "model confident in predicting a winner and cause it to stop predicting draws "
        "entirely, which collapses the macro F1 score. Our World Cup-only ELO system, "
        'which updates only every four years, creates what we term "productive '
        'imprecision": when many teams enter a tournament with similar ratings, the model '
        "treats draws as plausible outcomes. This finding relates to Dixon and Coles' "
        "(1997) observation about the difficulty of draw prediction and suggests that some "
        "degree of feature imprecision can actually be beneficial in small-sample "
        "classification."
    ),
    (
        "**Domain mismatch in expanded training.** We attempted to expand our training "
        "from 900 World Cup matches to 28,000 international matches, expecting that more "
        'data would help. Instead, every configuration performed worse. The reason is '
        'a fundamental domain mismatch: in qualifying matches and friendlies, the "home '
        'team" plays at their actual stadium with a roughly 60% win rate, but in World '
        'Cup matches, "home team" is a FIFA administrative label at a neutral venue with '
        "a roughly 50% win rate. The model learned a home-win bias from non-World Cup "
        "data that does not apply to the tournament. This illustrates an important "
        "lesson: more data is not always better when the additional data comes from a "
        "different distribution."
    ),
    (
        "**Temporal cross-validation as essential methodology.** Every model we tested "
        "showed 4 to 9 percentage points lower accuracy under temporal walk-forward "
        "cross-validation compared to standard stratified cross-validation. This gap "
        "confirms that shuffled cross-validation creates temporal information leakage "
        "in sports prediction. We recommend temporal cross-validation as the primary "
        "evaluation method for any time-series or sports prediction problem."
    ),
    (
        "**Unsupervised analysis.** PCA shows that roughly 15 of 37 principal components "
        "explain 90% of variance, indicating that our features capture diverse, "
        "non-redundant information. K-Means clustering with k = 3 produced an Adjusted "
        "Rand Index of only 0.022 against actual outcome labels, confirming that match "
        "outcomes do not form easily separable clusters in feature space. This underscores "
        "the inherent difficulty of the prediction task."
    ),
]

RECOMMENDATIONS = [
    (
        "Based on our findings, we offer several recommendations and directions for "
        "future work."
    ),
    (
        "**1. Incorporate betting odds as a feature or calibration target.** Pre-match "
        "betting odds from the sports betting market aggregate large amounts of public "
        "information about team form, injuries, and tactical matchups. Published research "
        "(Groll et al., 2019) has shown that odds-based features are among the strongest "
        "predictors for football match outcomes. Adding this as a feature could push "
        "accuracy beyond what ELO-based features alone can achieve."
    ),
    (
        "**2. Explore ordinal regression.** Football outcomes have a natural ordering "
        "(away win, draw, home win), and our current models treat them as unordered "
        "categories. Methods that exploit this ordinal structure, such as the approach by "
        "Hubacek et al. (2019), might improve draw prediction by treating it as an "
        "intermediate outcome rather than an independent class."
    ),
    (
        "**3. Prioritize regularization before feature engineering on small datasets.** "
        "Our finding that regularization outperforms feature engineering by a factor of "
        "four has practical implications beyond this project. For any prediction problem "
        "with limited training data, practitioners should first tune model complexity "
        "(minimum leaf size, maximum depth, or regularization penalties) before investing "
        "effort in gathering additional feature sources. Our experiments suggest that the "
        '"feature budget" on 900 rows is roughly 37 to 39 features, and going beyond '
        "that threshold causes more harm than benefit through overfitting."
    ),
    (
        "**4. Revisit expected goals with broader coverage.** Among the five external "
        "feature sources we tested, StatsBomb's expected goals had the lowest correlation "
        "with ELO (r = 0.35) and captured genuinely novel information about match quality "
        "rather than team strength. It failed in our experiments solely due to limited "
        "coverage (only 10% of training data had real values). If event-level data becomes "
        "available for additional World Cups beyond 2018 and 2022, this feature could be "
        "a valuable addition."
    ),
    (
        "**5. Develop a more sophisticated ELO integration.** Our World Cup-only ELO "
        "creates productive imprecision that helps draw prediction, but computing ELO "
        "from all international matches with a variable K-factor (weighting World Cup "
        "matches more heavily) could potentially capture both benefits: the richer update "
        "frequency of international data and the domain specificity of World Cup matches. "
        "This would require careful handling of the domain mismatch issue we identified."
    ),
    (
        "**Broader implications.** Our project shows that in applied prediction tasks "
        "with small datasets, model regularization and domain-specific evaluation "
        'methodology matter more than sophisticated features or complex algorithms. The '
        '"overconfidence" pattern we identified, in which precise features hurt '
        "minority-class prediction, may apply to other imbalanced classification tasks "
        "in domains such as medical diagnosis or fraud detection, where minority classes "
        "are both rare and practically important."
    ),
]

GITHUB_LINK = [
    "https://github.com/drdiguglielmo-personal/datascienceproject"
]

REFERENCES = [
    (
        "Dixon, M. J., and Coles, S. G. (1997). Modelling association football scores "
        "and inefficiencies in the football betting market. Journal of the Royal "
        "Statistical Society: Series C (Applied Statistics), 46(2), 265-280."
    ),
    (
        "Fjelstul, J. C. (2021). The Fjelstul World Cup Database v1.2.0. "
        "https://github.com/jfjelstul/worldcup"
    ),
    (
        "Groll, A., Ley, C., Schauberger, G., and Lock, H. (2019). A hybrid random "
        "forest to predict soccer matches in national and international tournaments. "
        "Journal of Quantitative Analysis in Sports, 15(4), 271-287."
    ),
    (
        "Hubacek, O., Sourek, G., and Zelezny, F. (2019). Exploiting sports-betting "
        "market using machine learning. International Journal of Forecasting, 35(2), "
        "783-796."
    ),
    (
        "Hvattum, L. M., and Arntzen, H. (2010). Using ELO ratings for match result "
        "prediction in association football. International Journal of Forecasting, 26(3), "
        "460-470."
    ),
    (
        "Maher, M. J. (1982). Modelling association football scores. Statistica "
        "Neerlandica, 36(3), 109-118."
    ),
]


def main():
    doc = Document(str(TEMPLATE))
    tables = doc.tables

    # Table 2: Executive Summary
    fill_cell(tables[2].rows[0].cells[0], EXEC_SUMMARY)

    # Table 3: Problem Definition
    fill_cell(tables[3].rows[0].cells[0], PROBLEM_DEF)

    # Table 4: Related Work
    fill_cell(tables[4].rows[0].cells[0], RELATED_WORK)

    # Table 5: Data Collection
    fill_cell(tables[5].rows[0].cells[0], DATA_COLLECTION)

    # Table 6: Data Preprocessing
    fill_cell(tables[6].rows[0].cells[0], DATA_PREPROCESSING)

    # Table 7: Data Documentation
    fill_cell(
        tables[7].rows[0].cells[0],
        DATA_DOCUMENTATION,
        figures=[("01_class_distribution.png", 4.5)],
    )

    # Table 8: Variables
    fill_cell(tables[8].rows[0].cells[0], VARIABLES)

    # Table 9: Algorithm 1 (Random Forest)
    fill_cell(tables[9].rows[0].cells[0], ALGO_1_RF)

    # Table 10: Algorithm 2 (XGBoost)
    fill_cell(tables[10].rows[0].cells[0], ALGO_2_XGB)

    # Table 11: Algorithm 3 (Other five)
    fill_cell(tables[11].rows[0].cells[0], ALGO_3_OTHERS)

    # Table 12: Baseline
    fill_cell(tables[12].rows[0].cells[0], BASELINE)

    # Table 13: Metrics
    fill_cell(
        tables[13].rows[0].cells[0],
        METRICS,
        figures=[
            ("03_model_comparison.png", 5.0),
            ("05_confusion_matrix_best.png", 4.0),
        ],
    )

    # Table 14: Discussion
    fill_cell(
        tables[14].rows[0].cells[0],
        DISCUSSION,
        figures=[
            ("06_feature_importance_rf.png", 4.5),
            ("04_temporal_vs_stratified_cv.png", 4.5),
        ],
    )

    # Table 15: Recommendations
    fill_cell(tables[15].rows[0].cells[0], RECOMMENDATIONS)

    # Table 16: GitHub link
    fill_cell(tables[16].rows[0].cells[0], GITHUB_LINK)

    # --- References at end of document ---
    doc.add_paragraph()  # spacing
    ref_heading = doc.add_paragraph("References")
    for run in ref_heading.runs:
        set_run_font(run, size=Pt(14), bold=True)
    ref_heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for ref_text in REFERENCES:
        p = doc.add_paragraph()
        run = p.add_run(ref_text)
        set_run_font(run)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(6)

    doc.save(str(OUTPUT))
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    main()
