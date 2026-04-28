"""
Generate the final Part 2 Report DOCX from docs/project_part2.md content,
following the structure of Part2_Report/CSE40467 - Project Part 2 Report Template.docx.

Output: Part2_Report/CSE40467 - Project Part 2 Report.docx (overwritten).

Style follows the template instructions:
  - Times New Roman, 11 pt, single line spacing
  - Boxes around section content (light gray border)
  - Italic prompts above each content box
  - Figure placeholders use [figures/<name>.png] so the team can swap manually.
"""

from __future__ import annotations

import pathlib
import re

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm

PROJECT_ROOT = pathlib.Path(__file__).resolve().parents[1]
OUT_PATH = PROJECT_ROOT / "Part2_Report" / "CSE40467 - Project Part 2 Report.docx"

FONT_NAME = "Times New Roman"
FONT_PT = 11


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def set_run_font(run, name=FONT_NAME, pt=FONT_PT, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(pt)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(k), name)


def set_paragraph_spacing(paragraph, line=1.0, before_pt=0, after_pt=4):
    pf = paragraph.paragraph_format
    pf.line_spacing = line
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)


def set_cell_borders(cell, color="999999", size="6"):
    """Light gray box around a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def add_paragraph_with_runs(parent, text_segments, *, bold=False, italic=False, alignment=None):
    """parent can be Document or a cell. text_segments is a list of (text, {bold, italic}) tuples."""
    p = parent.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    set_paragraph_spacing(p)
    for seg in text_segments:
        if isinstance(seg, str):
            run = p.add_run(seg)
            set_run_font(run, bold=bold, italic=italic)
        else:
            text, opts = seg
            run = p.add_run(text)
            set_run_font(
                run,
                bold=opts.get("bold", bold),
                italic=opts.get("italic", italic),
            )
    return p


def add_simple_paragraph(parent, text, *, bold=False, italic=False, pt=FONT_PT, alignment=None):
    p = parent.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    set_paragraph_spacing(p)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic, pt=pt)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=12, after_pt=6)
    run = p.add_run(text)
    set_run_font(run, bold=True, pt=14 if level == 1 else 12)
    return p


def add_italic_prompt(doc, text):
    """Italic prompt line above each content box (matches template style)."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=6, after_pt=2)
    run = p.add_run(text)
    set_run_font(run, italic=True, pt=FONT_PT)
    return p


# Markdown-style inline parser: handles **bold** and *italic* (and `code`).
INLINE_PATTERN = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")


def render_inline(paragraph, text, *, base_bold=False, base_italic=False):
    """Render a string with **bold**, *italic*, `code` runs into a paragraph."""
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, bold=base_bold, italic=base_italic)
        token = m.group(1)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True, italic=base_italic)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, bold=base_bold, italic=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, bold=base_bold, italic=base_italic, name="Courier New")
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, bold=base_bold, italic=base_italic)


def add_box_paragraph(doc, text, *, italic=False):
    """Add a single paragraph wrapped inside a 1x1 bordered table (the 'box')."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    # Remove default paragraph that python-docx adds in cells
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing(p)
    render_inline(p, text, base_italic=italic)
    add_spacer(doc)
    return table


def add_box_with_paragraphs(doc, paragraphs):
    """A single 1x1 box that contains multiple paragraphs / lines."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.paragraphs[0].text = ""
    first = True
    for line, kind in paragraphs:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        set_paragraph_spacing(p)
        if kind == "bullet":
            p.style = "List Bullet"
            render_inline(p, line)
        elif kind == "ordered":
            p.style = "List Number"
            render_inline(p, line)
        elif kind == "heading":
            run = p.add_run(line)
            set_run_font(run, bold=True)
        elif kind == "code":
            run = p.add_run(line)
            set_run_font(run, name="Courier New", pt=10)
        else:
            render_inline(p, line)
    add_spacer(doc)
    return table


def add_table_in_box(doc, headers, rows, *, header_bold=True):
    """Add a data table inside a bordered shell paragraph (just adds the table, no outer box)."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    # header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_borders(cell, color="666666")
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_spacing(p)
        run = p.add_run(str(h))
        set_run_font(run, bold=header_bold)
    # body rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            set_cell_borders(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p)
            render_inline(p, str(val))
    add_spacer(doc)
    return table


def add_figure_placeholder(doc, filename, caption=None):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=4, after_pt=4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[{filename}]")
    set_run_font(run, bold=True, color=RGBColor(0x80, 0x80, 0x80))
    if caption:
        cp = doc.add_paragraph()
        set_paragraph_spacing(cp, before_pt=0, after_pt=6)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        set_run_font(run, italic=True, pt=10)


def add_spacer(doc, pt=4):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before_pt=0, after_pt=pt)
    return p


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------
def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_PT)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for k in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(k), FONT_NAME)


def build_header(doc):
    add_simple_paragraph(doc, "CSE40467 - Data Science", bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_simple_paragraph(doc, "Milestone 1", italic=False)
    add_simple_paragraph(doc, "Prof. Diego Gomez-Zara", italic=True, pt=10)
    add_simple_paragraph(doc, "TAs: Nandini Banerjee and Colby Nelson", italic=True, pt=10)
    add_spacer(doc, pt=8)

    # Team info table
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info = [
        ("Team Number", "Group 3"),
        ("Members", "Dan Yoo, Drew DiGuglielmo, Michael Sorenson, Charlie Devine"),
        ("Date", "09 April 2026 (revised 28 April 2026 after Part 2 presentation feedback)"),
    ]
    for ri, (k, v) in enumerate(info):
        for ci, txt in enumerate((k, v)):
            cell = table.rows[ri].cells[ci]
            set_cell_borders(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p)
            run = p.add_run(txt)
            set_run_font(run, bold=(ci == 0))
    add_spacer(doc, pt=8)

    # Instructions
    add_simple_paragraph(doc, "Instructions:", bold=True)
    instructions = [
        "Maximum length: 20 pages, excluding members' contributions and references.",
        "Put your responses inside the boxes.",
        "You can add figures in the boxes.",
        "Do not alter the font size (11-point), line spacing (1-line), or font typography (Times New Roman).",
        "References should be added by the end of the document. Use the format that you prefer (e.g., APA, IEEE, ACM, etc.)",
        "Please do not forget to cite when you make claims or references.",
    ]
    for line in instructions:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p)
        run = p.add_run(line)
        set_run_font(run)
    add_spacer(doc, pt=8)


# ---------------------------------------------------------------------------
# Section content (text taken from docs/project_part2.md, structured for boxes)
# ---------------------------------------------------------------------------
def section_1_introduction(doc):
    add_heading(doc, "1. Introduction", level=1)

    add_italic_prompt(
        doc,
        "1.1. Executive Summary. A good executive summary is a concise and clear "
        "overview of a larger document that highlights its key points, findings, and "
        "recommendations, aimed at giving readers a quick and comprehensive understanding "
        "of its content.",
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "We predict FIFA Men's World Cup match outcomes from pre-match information only, "
                "framed as a three-class classification problem (home team win, away team win, or draw). "
                "The dataset has 964 matches across 22 tournaments from 1930 to 2022, with 900 matches used "
                "for training (1930 to 2018) and 64 matches held out as the test set (2022 Qatar World Cup). "
                "We engineered 39 leakage-safe features that span team strength proxies (ELO ratings, "
                "expanding-window win rates), short-term form (rolling five-match aggregates), match context "
                "(group vs knockout stage, host status), and head-to-head history. All features are computed "
                "with a date-level shift so that same-day match results never leak into a row's inputs.",
                "para",
            ),
            (
                "We compared seven supervised algorithms (KNN, Decision Tree, Naive Bayes, SVM with RBF kernel, "
                "Random Forest, MLP, and XGBoost) under temporal walk-forward cross-validation, where each fold "
                "trains on tournaments up to year T and validates on year T+1. The tuned Random Forest "
                "(300 trees, max_features='sqrt', min_samples_leaf=5, class_weight='balanced') reached 62.5% "
                "test accuracy and 0.578 macro F1 on the 2022 World Cup, which beats the 50% majority-class "
                "baseline by 12.5 percentage points and lands at the high end of the 53% to 55% range published "
                "by Hvattum and Arntzen (2010) for similar three-class tasks.",
                "para",
            ),
            (
                "Two findings are worth flagging up front. First, on a 900-row dataset, regularization beats "
                "feature engineering by roughly four to one. Increasing min_samples_leaf from 1 to 5 added 4.3 "
                "points of CV F1, while the best new feature (continent-advantage flags) added 1.0 point. "
                "Second, more precise team-strength features (FIFA points, squad market values) hurt performance "
                "because they made the model overconfident about the favorite and stopped it from predicting draws "
                "at all.",
                "para",
            ),
            (
                "After our Part 2 presentation, the reviewer (Prof. Gomez-Zara) raised two follow-up questions "
                "that we addressed and added to this report. (1) Because draws are the hardest class, what happens "
                "if you train a binary model on win-or-loss only, especially on knockout matches where draws are "
                "decided by penalties? Result: a binary RF on the same 39 features and the same 2022 test set "
                "(n=54 after dropping draws) reaches 70.4% accuracy and 0.724 F1. The knockout-only variant "
                "reaches 0.821 AUC on n=16 matches. (2) \"Home team\" at a neutral World Cup venue is just FIFA's "
                "team_1 label, so saying \"home wins are predicted best\" is not actionable. Re-labeling outcomes "
                "as favorite (higher pre-match ELO), draw, or upset reveals that the same Random Forest predicts "
                "favorite-wins with 79% recall but catches only 25% of upsets. The 2022 tournament happened to "
                "contain seven blowout-magnitude upsets (for example, Argentina-Saudi Arabia, ELO gap 250, and "
                "Cameroon-Brazil, gap 328), all of which the model rated as favorite-wins.",
                "para",
            ),
            (
                "Our recommendation, both for the course project and for similar small-data sports prediction "
                "work, is to keep temporal cross-validation as the default, prefer regularized ensembles over "
                "deeper models, and report results in a domain-meaningful framing (favorite vs upset, ELO tier "
                "matchups, confederation matchups) rather than the dataset's administrative home-vs-away labels.",
                "para",
            ),
        ],
    )

    add_italic_prompt(
        doc,
        "1.2. Problem definition. Explain what this problem is about. Discuss why it should be studied/analyzed.",
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "Predicting football match outcomes is a long-studied problem in sports analytics that remains "
                "genuinely hard. We chose the FIFA Men's World Cup version of this problem for four reasons.",
                "para",
            ),
            (
                "First, the World Cup is a closed-population, low-volume, high-stakes setting that exposes "
                "interesting statistical issues. There are 22 tournaments and roughly 64 matches per modern "
                "edition, so the entire historical dataset is small. The teams that appear are not a random "
                "sample of national teams; they are the ones that won qualification. This makes the prediction "
                "problem both interesting (selection effects, era drift) and methodologically demanding "
                "(nine-hundred rows is not a lot for tree-based ensembles).",
                "para",
            ),
            (
                "Second, football is a low-scoring, high-variance sport. A single goal often decides a match "
                "regardless of which side dominated possession or expected goals. This puts a hard ceiling on "
                "prediction accuracy. Published baselines for three-class international football outcome "
                "prediction sit around 53% to 55% on much larger datasets (Hvattum and Arntzen, 2010), which "
                "means even a few percentage points of improvement is meaningful and any claim above 65% should "
                "be regarded with suspicion.",
                "para",
            ),
            (
                "Third, the three-class structure is harder than binary classification because draws are "
                "intrinsically difficult to predict. Dixon and Coles (1997) noted decades ago that draws emerge "
                "from many distinct match dynamics: tactical caution in group stage matches when both sides "
                "need a point, evenly matched teams, and late-tournament fatigue. Our experiments reproduce "
                "this finding: every single model we trained had its lowest per-class recall on draws.",
                "para",
            ),
            (
                "Fourth, the dataset has an unusual labeling quirk that makes evaluation interpretation tricky. "
                "The \"home team\" in a World Cup is not the team playing at their home stadium (the venue is "
                "neutral, except in the rare case of a host nation match). It is a FIFA-assigned team_1 label, "
                "and the historical home-team win rate of about 57% reflects FIFA seeding conventions rather "
                "than any genuine home-field effect. We came back to this issue after presentation feedback "
                "(see Section 5.3) and re-framed our results so that they read more naturally and answer the "
                "practical question of who to bet on.",
                "para",
            ),
            (
                "We believe this problem is worth studying because it sits at an instructive junction in applied "
                "data science. Small dataset, time-ordered observations, imbalanced classes, an attractive but "
                "partly misleading label, and a domain (sports) where careful evaluation methodology matters as "
                "much as model choice. Many real applied tasks in fraud detection, medical diagnosis, and "
                "operations forecasting share at least three of those four properties.",
                "para",
            ),
        ],
    )


def section_2_related_work(doc):
    add_heading(doc, "2. Related Work", level=1)
    add_italic_prompt(doc, "2.1. Describe related work with this data source.")
    add_box_with_paragraphs(
        doc,
        [
            (
                "Football outcome prediction has been an active research area since the 1980s. We organize the "
                "relevant literature into four threads, each of which informed a specific choice in our pipeline.",
                "para",
            ),
            ("Score-based generative models.", "heading"),
            (
                "Maher (1982) introduced one of the earliest statistical models, treating goal counts as "
                "independent Poisson processes with team-specific attack and defense parameters. The decomposition "
                "into offensive and defensive components influenced our feature design: we have separate "
                "home_hist_goals_per_game and home_hist_goals_conceded_per_game features, plus their interaction "
                "(home_attack_x_away_defense). Dixon and Coles (1997) extended the Poisson model with a low-score "
                "correlation correction that improved draw prediction. Their core observation, that draws are "
                "systematically under-predicted by classifiers trained on imbalanced data, is one we hit "
                "repeatedly throughout our experiments.",
                "para",
            ),
            ("ELO-based discriminative models.", "heading"),
            (
                "Hvattum and Arntzen (2010) ran a systematic evaluation of ELO ratings as predictors for "
                "international football, on roughly 10,000 matches. Their best model, an ELO-difference-based "
                "ordinal logistic regression, reached around 53% to 55% three-class accuracy. Two findings "
                "transferred directly into our project. First, ELO difference is the single strongest pre-match "
                "feature, which our Random Forest feature importance analysis confirms (elo_diff is the top "
                "feature in every configuration we tested). Second, even with substantially more data, the "
                "ceiling is around 55% on three-class outcomes, which calibrated our expectations about "
                "plausible test-set numbers.",
                "para",
            ),
            ("Ordinal and market-aware approaches.", "heading"),
            (
                "Hubacek, Sourek, and Zelezny (2019) developed a score-based ordinal framework for the 2018 "
                "Soccer Prediction Challenge. They argued that the natural ordering (away win < draw < home win) "
                "carries information that standard multiclass classifiers ignore. Our models do treat the three "
                "classes as unordered, so an ordinal extension is a clear future-work direction. We do, however, "
                "use the Ranked Probability Score (Constantinou and Fenton, 2012) as one of our metrics, which "
                "respects the ordinal structure during evaluation.",
                "para",
            ),
            ("Tournament-style ensembles and feature studies.", "heading"),
            (
                "Groll, Ley, Schauberger, and Lock (2019) organized a 26-team prediction tournament for the 2018 "
                "World Cup. The headline finding from that competition was that no single algorithm dominates; "
                "thoughtful feature engineering and honest evaluation matter more than algorithmic choice. This "
                "is consistent with our results, where a tuned Random Forest beats XGBoost on this dataset size, "
                "and where the largest single improvement came from a regularization parameter rather than a "
                "clever model architecture. Tax and Joustra (2015) compared Naive Bayes, logistic regression, "
                "and random forests on Dutch league data and found that ensembles tend to win when feature "
                "inputs are heterogeneous, which our seven-model comparison echoes.",
                "para",
            ),
            ("Our positioning.", "heading"),
            (
                "We see our contribution as four-fold. (1) We document strict temporal evaluation on a small "
                "World-Cup-only dataset, with quantification of how much standard k-fold CV overestimates "
                "performance. (2) We engineered a leakage-safe feature pipeline whose construction details "
                "(date-level shift, fold-respecting cold starts) are reproducible. (3) We tested five external "
                "feature sources, with both positive and negative results, and explained mechanistically why "
                "some hurt the model (the \"overconfidence\" pattern). (4) After presentation feedback, we added "
                "a binary no-draw variant and an ELO-tier-based result reframing that together change how we "
                "describe what the model is doing.",
                "para",
            ),
        ],
    )


def section_3_data(doc):
    add_heading(doc, "3. Data Description", level=1)

    # 3.1
    add_italic_prompt(
        doc,
        "3.1. Data Collection. Explain how data collection was conducted. You should explain its "
        "origin, context, and purpose, and clearly describe the data source as well as how the data "
        "was collected or generated. If your dataset comes from a website, describe the source and "
        "explain how the original researchers potentially obtained the data.",
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "Our primary data source is the Fjelstul World Cup Database v1.2.0 (Fjelstul, 2021), a publicly "
                "available open-source repository of FIFA World Cup matches and tournaments. We use two raw CSV "
                "files: matches.csv, which contains every World Cup match record (men's and women's combined), "
                "and tournaments.csv, which provides tournament-level metadata such as the year, host nation, "
                "and tournament identifier. Both files were downloaded directly from the GitHub repository at "
                "https://github.com/jfjelstul/worldcup. No web scraping or API calls were needed at runtime.",
                "para",
            ),
            (
                "The original compiler (Joshua Fjelstul) appears to have assembled the records from official "
                "FIFA archives and historical competition reports, structuring them into normalized relational "
                "tables. After filtering to men's World Cup matches only (we match on the tournament name field "
                "containing \"FIFA Men's World Cup\"), the dataset contains 964 matches across 22 tournaments "
                "from 1930 to 2022, featuring 84 distinct national teams. Each match record includes 38 columns "
                "covering tournament and match identifiers, stage information (group stage indicator, knockout "
                "stage indicator, named round), home and away team names and codes, scores, extra-time and "
                "penalty indicators, venue details (stadium, city, country), and outcome labels.",
                "para",
            ),
            (
                "For our feature experimentation phase (Phase 3 in our internal terminology), we additionally "
                "integrated five external sources, all publicly available:",
                "para",
            ),
            ("International football results (49,287 matches, 1872 to 2026), from a Kaggle dataset by Mart Jurisoo (CC0).", "ordered"),
            ("FIFA world rankings (67,894 records, 1992 to 2024), from the Dato-Futbol GitHub repository.", "ordered"),
            ("Transfermarkt player valuations (47,702 player records and roughly 616,000 historical valuations, 2004 to 2026).", "ordered"),
            ("World Cup squad rosters (13,843 entries, 1930 to 2022), from a companion Fjelstul repository.", "ordered"),
            ("StatsBomb expected-goals statistics (256 team-match records, 2018 and 2022 World Cups only).", "ordered"),
            (
                "All of these are documented in data_clean/ with their source URLs and licenses. Note that we "
                "ultimately rejected most of them from the final feature set (see Section 6.1 for why); we keep "
                "them in the repository because the rejection itself is a finding.",
                "para",
            ),
            (
                "Selection caveat. Rows are restricted to teams that qualified for a World Cup. This is a "
                "non-random selection: weaker national teams systematically do not appear, and within tournaments "
                "certain confederations are over-represented (UEFA and CONMEBOL together account for the majority "
                "of teams in most editions). Our results should therefore be read as conditional on tournament "
                "qualification, not as a model of all national-team football.",
                "para",
            ),
        ],
    )

    # 3.2
    add_italic_prompt(
        doc,
        "3.2. Data Preprocessing. Explain all the steps that you performed to transform the raw data "
        "into a manageable dataset (e.g., merge, impute missing values, standardize, normalize). If "
        "your data did not require any data pre-processing, please discuss any potential quality "
        "issues of the data (e.g., how representative is the dataset? Who pre-processed the dataset?)",
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "The pipeline has two main stages, plus a small post-processing step for the test-time analyses "
                "we added in Section 5.3.",
                "para",
            ),
            ("Cleaning (scripts/clean_worldcup.py, 173 lines).", "heading"),
            ("Filter the raw matches table to men's World Cups by matching on tournament_name containing the literal string \"FIFA Men's World Cup\".", "ordered"),
            ("Restrict matches to those tournaments via tournament_id, then merge tournament_year onto each match.", "ordered"),
            (
                "Type cleaning. Date columns are parsed with pd.to_datetime(errors='coerce') to safely handle "
                "malformed entries. Boolean-like columns (group_stage, knockout_stage, extra_time, "
                "penalty_shootout, win and draw flags) are coerced to integer 0/1 with "
                "pd.to_numeric(...).fillna(0). Score columns are ensured numeric.",
                "ordered",
            ),
            (
                "Train/test split by year: train = 1930 to 2018 (900 matches), test = 2022 (64 matches). This "
                "split simulates the deployment setting we care about, which is \"given everything up to today's "
                "tournament, predict the next one.\"",
                "ordered",
            ),
            ("Feature engineering (scripts/feature_engineering.py, 1,985 lines).", "heading"),
            (
                "The training and test sets are concatenated and sorted chronologically before any rolling or "
                "expanding-window computation, so that 2022 features depend only on 1930 to 2018 history at "
                "construction time. We then split back into train and test files.",
                "ordered",
            ),
            (
                "Leakage prevention. Expanding-window aggregates and rolling-window form features are computed "
                "with a date-level shift of 1, not a row-level shift of 1. World Cup days frequently host "
                "multiple matches; with a row-level shift, two same-day matches could leak each other's results "
                "into each other's features. The date-level shift treats all same-day matches as if they had "
                "been played on the same instant, which is the conservative choice.",
                "ordered",
            ),
            (
                "Post-match column exclusion. Score columns, margin columns, win flags, and penalty strings are "
                "not allowed as model inputs. They may inform target construction or historical updates of "
                "expanding stats, but they are dropped before the feature CSV is written.",
                "ordered",
            ),
            (
                "Cold-start imputation. For teams in their first World Cup match in the dataset, win/draw/loss "
                "rates are filled with 0.33 (uniform prior), counts with 0, ELOs at 1500 (the standard initial "
                "value), rest days at the median of available values, and interaction features at 0. The exact "
                "fills are documented in code and in README.md Section 4.3.",
                "ordered",
            ),
            ("Modeling preprocessing (Part2_Models_and_Results.ipynb).", "heading"),
            (
                "We drop identifier and label columns from the feature matrix (match_date, year, "
                "home_team_name, away_team_name, result).",
                "ordered",
            ),
            (
                "We exclude experimental feature groups (FIFA points, qualifying records, squad value, xG, "
                "international-history features) that we tested and rejected. The final feature count is 39.",
                "ordered",
            ),
            (
                "StandardScaler is fit on training features only and used to transform both train and test, "
                "which prevents test-set leakage through the scaler.",
                "ordered",
            ),
            (
                "LabelEncoder is applied to the target for sklearn compatibility, especially for the "
                "MLPClassifier with early stopping.",
                "ordered",
            ),
            ("Quality notes.", "heading"),
            (
                "The data is not a random sample of all football. It is elite-tournament matches with "
                "longitudinal drift across 92 years (rule changes, expansion to 32 teams, format adjustments). "
                "\"Home team\" is an administrative label at usually neutral venues (the host nation is one "
                "common exception). External data sources have partial coverage (FIFA rankings: 60%, squad "
                "value: 30%, xG: 10%) and we either imputed or dropped them depending on whether they helped "
                "temporal CV metrics.",
                "para",
            ),
        ],
    )

    # 3.3
    add_italic_prompt(
        doc,
        "3.3. Data Documentation. Please include a clear and comprehensive description of your dataset "
        "with descriptive statistics. In addition, explicitly state any assumptions, limitations, and "
        "preprocessing decisions you made. Your documentation should provide enough detail for someone "
        "else to fully understand the dataset and reproduce your work.",
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "The cleaned dataset has 964 matches with 38 raw columns each. Training: 900 matches across 21 "
                "tournaments (1930 to 2018). Test: 64 matches in the 2022 Qatar World Cup. There are 84 distinct "
                "national teams. The full codebook lives at docs/worldcup_subset_codebook.csv.",
                "para",
            ),
            ("Target distribution.", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["Class", "Train Count", "Train %", "Test Count", "Test %"],
        [
            ["Home team win", "513", "57.0", "32", "50.0"],
            ["Away team win", "218", "24.2", "16", "25.0"],
            ["Draw", "169", "18.8", "16", "25.0"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "The shift between train and test is itself informative. Training is heavily home-skewed, but "
                "the 2022 test set is more balanced. A model that exploits the home-win prior too aggressively "
                "will be penalized on 2022.",
                "para",
            ),
            ("Dataset properties.", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["Property", "Value"],
        [
            ["Total matches", "964"],
            ["Training matches", "900"],
            ["Test matches", "64"],
            ["Distinct teams", "84"],
            ["Final modeling features", "39"],
            ["Goals per match (home, mean)", "1.78"],
            ["Goals per match (away, mean)", "1.05"],
            ["Null values after cleaning", "0"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            ("Assumptions.", "heading"),
            ("Pre-match information in the database is truthfully recorded.", "bullet"),
            (
                "World-Cup-only ELO is a deliberate modeling choice. It updates only on World Cup matches "
                "(so a four-year gap between updates), which is less precise than ELO computed from all "
                "international matches but preserves draw plausibility (the \"productive imprecision\" we "
                "describe in Section 6.1).",
                "bullet",
            ),
            (
                "The train/test split by year assumes the deployment setting is \"predict the next World Cup,\" "
                "not \"predict random held-out matches across all eras.\"",
                "bullet",
            ),
            ("Limitations.", "heading"),
            (
                "Tiny test set (n=64) means single-tournament variance dominates. As we report in Section 5.1, "
                "single-seed Random Forest accuracy fluctuates between 0.594 and 0.625 across random seeds 0 to "
                "9 with the same data and hyperparameters.",
                "bullet",
            ),
            (
                "The 92-year span means the oldest matches may be only weakly comparable to modern football "
                "(substitutions did not exist before 1970, the back-pass rule was introduced in 1992, etc.).",
                "bullet",
            ),
            ("Qualification selection: there are no non-qualifying national teams in our data.", "bullet"),
            (
                "The \"home team\" label is a FIFA administrative artifact, not a true home advantage, except "
                "for the host's own matches.",
                "bullet",
            ),
            ("Reproducibility.", "heading"),
            (
                "All transformations are deterministic given a fixed random seed for the model fits. The full "
                "pipeline is python3 scripts/clean_worldcup.py && python3 scripts/feature_engineering.py, after "
                "which the two notebooks (Data Science Report.ipynb for Part 1 EDA and Part2_Models_and_Results.ipynb "
                "for Part 2 modeling) reproduce all figures. docs/worldcup_subset_codebook.csv documents raw columns, "
                "and README.md documents engineered features and the experiment log.",
                "para",
            ),
        ],
    )

    # 3.4
    add_italic_prompt(
        doc,
        "3.4. Variables. Describe which columns are features (independent variables) and which ones "
        "are your target/class/dependent variables.",
    )
    add_box_with_paragraphs(
        doc,
        [
            ("Target variable (dependent).", "heading"),
            (
                "result, a three-class categorical label taking values {home team win, away team win, draw}.",
                "para",
            ),
            ("Feature variables (independent).", "heading"),
            ("39 features, organized into 10 groups:", "para"),
            (
                "(A) Team historical performance (12 features). Cumulative win rate, draw rate, goals per game, "
                "goals conceded per game, and matches played for both teams, plus difference features for win "
                "rate and goals per game. These are expanding-window aggregates computed from the team's full "
                "World Cup history up to the previous match (date-level shift).",
                "para",
            ),
            (
                "(B) Match context (3 features). is_group_stage and is_knockout binary indicators, plus "
                "stage_ordinal (group stage = 0, round of 16 = 1, ..., final = 6). Stage matters because "
                "group-stage matches behave differently from knockouts, especially on draw rates.",
                "para",
            ),
            (
                "(C) Host advantage (2 features). Binary flags for whether each team is the tournament's host "
                "nation. Hosts win at meaningfully elevated rates (Brazil 1950, France 1998, South Korea 2002 "
                "are extreme cases).",
                "para",
            ),
            (
                "(D) World Cup experience (3 features). Count of distinct prior World Cups each team participated "
                "in, plus the difference. This is a stable proxy for tournament-level seasoning that ELO does "
                "not capture.",
                "para",
            ),
            (
                "(E) Head-to-head (5 features). Total prior World Cup meetings, home team wins in those meetings, "
                "away team wins, draws, and the home team's win rate (defaulted to 0.33 when no prior meetings "
                "exist).",
                "para",
            ),
            (
                "(F) ELO ratings (3 features). home_elo, away_elo, and elo_diff. ELO is a relative team-strength "
                "score that originated in chess (Elo, 1978): every team starts at 1500, points transfer between "
                "teams after each match, and the size of the transfer is governed by a K-factor (we use K=32). "
                "A team that beats a stronger opponent gains more points than one that beats a weaker opponent. "
                "We compute ELO from World Cup matches only, which means ratings update once every four years. "
                "We discuss the consequence of this choice in Sections 4.1 and 6.1. (This paragraph is a direct "
                "response to a presentation Q&A request to define ELO; the README §F has the same explanation "
                "for non-report readers.)",
                "para",
            ),
            (
                "(G) Rolling form (4 features). Win rate and goals per game over each team's last five World Cup "
                "matches. Falls back to expanding aggregates when a team has fewer than five prior matches.",
                "para",
            ),
            (
                "(H) Rest days (2 features). Days since each team's previous match, capped at 365 (a fresh slate "
                "for a new tournament).",
                "para",
            ),
            (
                "(I) Interaction features (3 features). home_attack_x_away_defense, away_attack_x_home_defense, "
                "and elo_x_form_diff. Captures multiplicative effects that linear models cannot represent natively.",
                "para",
            ),
            (
                "(J) Home continent advantage (2 features, the \"Phase 3\" addition we kept). Binary flags "
                "indicating whether each team's FIFA confederation matches the host country's confederation. The "
                "empirical pattern is that European teams win disproportionately at European World Cups (and "
                "similarly for South America), even after controlling for ELO. This was the only Phase 3 feature "
                "that improved temporal CV performance.",
                "para",
            ),
            (
                "The complete EXCLUDE list (FIFA points, qualifying records, squad value, xG, international ELO, "
                "international H2H, international rolling form) is in the notebook (EXCLUDE_FEATURES) and matches "
                "the script in scripts/binary_no_draw_model.py and scripts/test_set_breakdown.py. We tested those "
                "features and they hurt temporal CV performance, which we explain in Section 6.1 (the "
                "\"overconfidence\" pattern).",
                "para",
            ),
        ],
    )


def section_4_modeling(doc):
    add_heading(doc, "4. Modeling", level=1)

    # 4.1
    add_italic_prompt(
        doc,
        "4.1. Description of Algorithm 1 and its model. Please include details about the parameters. "
        "Did you fine-tune the parameters?",
    )
    add_box_with_paragraphs(
        doc,
        [
            ("Algorithm 1: Random Forest (best model).", "heading"),
            (
                "Random Forest is our primary model. It is a bagging-based ensemble that builds many decision "
                "trees, each on a bootstrap sample of the training data and using a random subset of features at "
                "each split, and averages their predicted class probabilities. We chose it for two reasons. First, "
                "bagging reduces variance, which is the dominant source of error on a 900-row dataset where any "
                "single tree overfits aggressively. Second, RF gives interpretable feature importance scores, "
                "which we used during feature selection.",
                "para",
            ),
            ("Final hyperparameters.", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["Parameter", "Value", "Rationale"],
        [
            ["n_estimators", "300", "Plateau in CV F1 beyond ~200 trees; 300 gives a small additional smoothing benefit"],
            ["max_features", "'sqrt'", "sqrt(39) ≈ 6 features per split, the standard tree-ensemble default"],
            ["min_samples_leaf", "5", "Most impactful single hyperparameter (see below)"],
            ["class_weight", "'balanced'", "Compensates for the 57/24/19 class imbalance in training"],
            ["random_state", "42 (also reported across seeds 0 to 9)", "See Section 5.1 for the seed-variance discussion"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            ("Hyperparameter tuning process.", "heading"),
            (
                "We grid-searched over n_estimators in {100, 200, 300, 500}, max_features in {0.2, 0.3, 'sqrt', "
                "'log2'}, and min_samples_leaf in {1, 3, 5} using 4-fold temporal walk-forward cross-validation. "
                "The most consequential parameter by a large margin was min_samples_leaf. Increasing it from 1 "
                "to 5 improved CV macro F1 from 0.488 to 0.531 (+0.043), the single largest gain anywhere in the "
                "project, exceeding the benefit of any feature addition we tested. The mechanism is "
                "straightforward regularization: requiring at least five samples per leaf prevents trees from "
                "carving out tiny subsets of the training data that memorize noise.",
                "para",
            ),
            ("Class imbalance handling.", "heading"),
            (
                "class_weight='balanced' assigns inversely proportional weights so that the minority draw class "
                "receives roughly three times the per-sample weight of the majority home-win class. Without it "
                "the model nearly stops predicting draws (we verified this with class_weight=None runs).",
                "para",
            ),
            ("Test-set results.", "heading"),
            (
                "Random Forest reached 0.625 accuracy and 0.578 macro F1 on the 2022 test set with "
                "random_state=42. Across 10 random seeds the accuracy varied from 0.594 to 0.625 with mean "
                "0.606, which we discuss honestly in Section 5.1.",
                "para",
            ),
            (
                "The top six features by Random Forest mean decrease in impurity are elo_diff, home_elo, "
                "away_elo, elo_x_form_diff, hist_win_rate_diff, and home_hist_win_rate. Four of the top six are "
                "ELO-derived, which echoes Hvattum and Arntzen (2010).",
                "para",
            ),
            ("[figures/06_feature_importance_rf.png]", "para"),
        ],
    )

    # 4.2
    add_italic_prompt(
        doc,
        "4.2. Description of Algorithm 2 and its model. Please include details about the parameters. "
        "Did you fine-tune the parameters?",
    )
    add_box_with_paragraphs(
        doc,
        [
            ("Algorithm 2: XGBoost (gradient boosting).", "heading"),
            (
                "XGBoost is our second model, chosen specifically because it represents a fundamentally different "
                "ensemble strategy from Random Forest. Random Forest is bagging (independent trees, averaged). "
                "XGBoost is boosting (trees built sequentially, each correcting the previous tree's residuals "
                "via gradient descent on a loss function). We wanted to test whether sequential error correction "
                "beats parallel averaging on this dataset size.",
                "para",
            ),
            ("Hyperparameters tested. We compared three configurations under temporal CV:", "heading"),
            ("max_depth=3, n_estimators=200, learning_rate=0.05", "bullet"),
            ("max_depth=4, n_estimators=150, learning_rate=0.08", "bullet"),
            ("max_depth=5, n_estimators=100, learning_rate=0.1", "bullet"),
            (
                "We added L1 and L2 regularization (reg_alpha=0.5, reg_lambda=2.0) to control overfitting, and we "
                "up-weighted draw rows by 1.5x via sample_weight to address the class imbalance. The depth-3, "
                "200-estimator, low-learning-rate configuration won under temporal CV.",
                "para",
            ),
            ("Results and comparison with Random Forest.", "heading"),
            (
                "XGBoost achieved 0.594 test accuracy and 0.514 macro F1, both below the tuned Random Forest. "
                "Temporal CV F1 was 0.434 (RF: 0.541). The gap is consistent across seeds.",
                "para",
            ),
            ("Why does Random Forest win on this data?", "heading"),
            (
                "Our reading is that boosting's advantage shows up on larger datasets where each sequential tree "
                "has enough signal to learn a meaningful residual pattern. With only 900 training rows, the "
                "residuals are dominated by noise rather than signal, so each subsequent tree fits noise and the "
                "overall ensemble overfits even with regularization. Random Forest's variance reduction by "
                "independent averaging is more robust at this scale. This matches the broader literature finding "
                "that bagging tends to outperform boosting on small datasets unless extremely careful tuning is "
                "applied.",
                "para",
            ),
        ],
    )

    # 4.3
    add_italic_prompt(
        doc,
        "4.3. Description of Algorithm 3 and its model. Please include details about the parameters. "
        "Did you fine-tune the parameters?",
    )
    add_box_with_paragraphs(
        doc,
        [
            ("Algorithm 3: five additional learners (KNN, Decision Tree, Naive Bayes, SVM, MLP).", "heading"),
            (
                "To meet the rubric requirement and to map out the algorithmic landscape, we evaluated five more "
                "models. Each one was tuned where applicable using temporal CV.",
                "para",
            ),
            (
                "K-Nearest Neighbors. Distance-based classifier. Tuned k over {3, 5, 7, 9}; k=7 won. Used "
                "distance-weighted voting and StandardScaler-normalized inputs (KNN is sensitive to scale). It "
                "is included as a non-parametric baseline that makes no assumptions about feature distributions. "
                "Result: 53.1% accuracy, 0.463 F1.",
                "para",
            ),
            (
                "Decision Tree. Single tree with max_depth tuned over {3, 5, 7, 10, None}; max_depth=5 won "
                "(deeper trees overfit). Mostly useful for interpretability and feature importance ranking, "
                "which independently confirmed elo_diff as the top feature. Result: 51.6% accuracy, 0.469 F1.",
                "para",
            ),
            (
                "Gaussian Naive Bayes. Probabilistic classifier that assumes features are conditionally "
                "independent given the class. We knew this assumption is violated by our correlated features "
                "(home_elo and home_hist_win_rate both encode team strength), and the model performs worst as "
                "expected. Result: 32.8% accuracy, 0.316 F1, below the majority baseline. This is included to "
                "demonstrate the cost of an inappropriate independence assumption.",
                "para",
            ),
            (
                "SVM with RBF kernel. Tuned C over {0.1, 1, 10, 100}; C=1 won. Used class_weight='balanced' and "
                "probability=True for RPS calculation. Result: 51.6% accuracy, 0.392 F1, but second-best RPS at "
                "0.139. The strong RPS suggests SVM's predicted probabilities are reasonably well-calibrated "
                "even when its argmax predictions are not the most accurate.",
                "para",
            ),
            (
                "Neural Network (MLPClassifier). Two hidden layers with 64 and 32 neurons, ReLU activations, "
                "Adam optimizer, early stopping, and a maximum of 500 iterations. Result: 54.7% accuracy, 0.511 "
                "F1. The MLP showed considerable instability across seeds, with temporal CV accuracy of only "
                "41.1% (the lowest of any model). On 900 training samples, neural networks struggle to learn "
                "stable patterns even with early stopping.",
                "para",
            ),
        ],
    )

    # 4.4 supplementary
    add_italic_prompt(
        doc,
        "4.4. Supplementary modeling: binary no-draw variants (presentation feedback). This subsection "
        "responds to a Part 2 presentation comment and is not a fourth required algorithm; it reuses "
        "the Random Forest from 4.1 in two reduced-task variants.",
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "After the Part 2 presentation, the reviewer suggested that since draws are the hardest class, "
                "we should try a binary win-or-loss version of the problem, especially on knockout-stage matches "
                "where draws are rare (in the modern format, knockouts are decided by extra time and penalty "
                "kicks, so most knockout records are not draws). We added two variants in "
                "scripts/binary_no_draw_model.py.",
                "para",
            ),
            (
                "Variant A: binary all-stages. Drop draws from training and test, predict 1 = home team win and "
                "0 = away team win on the remaining matches (731 train, 54 test in 2022). The temporal split is "
                "identical to the 3-class model: val_years=(2010, 2014, 2018), train = year < val_year, scaler "
                "fit only on the training fold. Critically, draws are filtered after splitting, not before, so "
                "no future-tournament information leaks into training. The script asserts this on every run.",
                "para",
            ),
            (
                "Variant B: binary knockout-only. Filter both train and test to knockout matches "
                "(is_knockout==1). World Cup knockouts are naturally binary because draws are decided by extra "
                "time and penalties. Only four knockout matches in the entire 1930 to 2018 training set are "
                "labeled \"draw\" in the dataset (all are 1934 to 1938 replayed matches), and zero of the sixteen "
                "2022 knockout matches are draws. Train n=226, test n=16.",
                "para",
            ),
            (
                "Hyperparameters. Same Random Forest configuration as the 3-class model (300 trees, sqrt features, "
                "min_samples_leaf=5, balanced class weights). We also report Logistic Regression with balanced "
                "class weights as a contrast.",
                "para",
            ),
            ("Results.", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["Variant", "Model", "n_train", "n_test", "Test Acc", "Test F1", "Test AUC"],
        [
            ["3-class baseline", "RF (single seed 42)", "900", "64", "0.625", "0.578", "n/a"],
            ["Binary all-stages", "LogReg", "731", "54", "0.537", "0.490", "0.562"],
            ["**Binary all-stages**", "**RF**", "**731**", "**54**", "**0.704**", "**0.724**", "**0.727**"],
            ["Binary knockout-only", "LogReg", "226", "16", "0.500", "0.556", "0.718"],
            ["**Binary knockout-only**", "**RF**", "**226**", "**16**", "**0.688**", "**0.762**", "**0.821**"],
            ["Slice (full-train, KO-only test)", "RF", "731", "16", "0.688", "0.762", "0.744"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            ("Reading the result.", "heading"),
            (
                "The binary all-stages RF gains +7.9 percentage points of accuracy and +0.146 macro F1 versus "
                "the 3-class RF on the same 39 features and the same temporal split. This confirms the "
                "reviewer's intuition: most of the 3-class error came from draws. The knockout-only variant "
                "pushes AUC to 0.821 but n=16 is fragile (one upset moves accuracy by 6.25 points), so the "
                "all-stages variant is the more robust headline.",
                "para",
            ),
            (
                "The Logistic Regression numbers are interesting in a negative way. With "
                "class_weight='balanced', LogReg over-corrects toward away wins given the natural 73/27 "
                "home/away imbalance in the no-draw dataset. The RF handles this far better, which is consistent "
                "with our reading throughout the project that bagging-based ensembles are the right tool at this "
                "dataset size.",
                "para",
            ),
            ("[figures/binary_no_draw_test_metrics.png]", "para"),
            ("[figures/binary_no_draw_roc.png]", "para"),
            ("[figures/binary_no_draw_knockout_confusion.png]", "para"),
        ],
    )


def section_5_evaluation(doc):
    add_heading(doc, "5. Evaluation", level=1)

    # 5.1
    add_italic_prompt(doc, "5.1. Baseline. Please describe your baseline for the evaluation.")
    add_box_with_paragraphs(
        doc,
        [
            (
                "Our baseline is the majority-class classifier: always predict home team win, the most frequent "
                "training label.",
                "para",
            ),
        ],
    )
    add_table_in_box(
        doc,
        ["Set", "Baseline Accuracy"],
        [
            ["Training (1930 to 2018, n=900)", "57.0%"],
            ["Test (2022 Qatar, n=64)", "50.0%"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "We chose this baseline over a uniform-random classifier (which would give 33% on three balanced "
                "classes) because it represents the simplest non-trivial strategy a human might use without any "
                "model: \"the FIFA-designated home team usually wins, so guess that.\" Any useful model must "
                "clear it. The tuned Random Forest beats it by 12.5 percentage points (62.5% vs 50.0%) on the "
                "test set, which we treat as evidence that our 39 features carry real predictive signal beyond "
                "the majority prior.",
                "para",
            ),
            (
                "We also use the published 53% to 55% three-class accuracy from Hvattum and Arntzen (2010) as an "
                "external comparison. Their dataset is much larger and broader (international football, not "
                "World-Cup-only), so the comparison is not apples-to-apples, but it sets a sensible expectation "
                "for what is achievable.",
                "para",
            ),
            ("A note on test-set seed sensitivity.", "heading"),
            (
                "With n=64, single-seed accuracy is genuinely noisy. Running the same Random Forest with "
                "random_state in {0, 1, 2, ..., 9}, accuracy ranged from 0.594 to 0.625 (mean 0.606, std 0.012, "
                "median around 0.609). The 0.625 we report from random_state=42 sits at the top of this range. "
                "This is not a bug; it is the expected behavior when one upset shifts accuracy by 1/64 ≈ 1.56 "
                "points. For the reframed slice analysis in Section 5.3 we therefore use a 10-seed ensemble with "
                "averaged class probabilities, which reduces noise per slice.",
                "para",
            ),
        ],
    )

    # 5.2
    add_italic_prompt(doc, "5.2. Metrics. Describe the metrics and their results.")
    add_box_with_paragraphs(
        doc,
        [
            (
                "For the three-class problem we report three metrics, all standard in the football-prediction "
                "literature.",
                "para",
            ),
            (
                "Accuracy. Fraction of correct predictions. Easy to interpret but can be misleading on "
                "imbalanced classes (a model that always predicts home can still look \"okay\" by accuracy).",
                "para",
            ),
            (
                "Macro F1. Unweighted average of per-class F1 scores. Penalizes models that ignore any class, "
                "which is exactly what we need given that draws are the most-ignored class. This is our primary "
                "metric for model comparison.",
                "para",
            ),
            (
                "Ranked Probability Score (RPS). Squared-error distance between predicted and actual cumulative "
                "class distributions, with classes ordered as away win < draw < home win. Lower is better. RPS "
                "is the standard metric in sports forecasting (Constantinou and Fenton, 2012) because it respects "
                "ordinal structure: predicting \"70% home win, 25% draw, 5% away win\" when the actual outcome "
                "is a draw is closer than predicting \"70% home win, 5% draw, 25% away win\", and RPS reflects "
                "this asymmetry while accuracy and F1 do not.",
                "para",
            ),
            (
                "For the binary no-draw experiments (Section 4.4) we additionally report ROC-AUC, since binary "
                "classification with calibrated probabilities is the natural framing.",
                "para",
            ),
            ("Results on 2022 test set, three-class problem.", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["Model", "Accuracy", "Macro F1", "RPS"],
        [
            ["**Random Forest (tuned)**", "**0.625**", "**0.578**", "**0.133**"],
            ["XGBoost", "0.594", "0.514", "0.146"],
            ["Neural Network (MLP)", "0.547", "0.511", "0.148"],
            ["SVM (RBF)", "0.563", "0.524", "0.139"],
            ["KNN", "0.531", "0.463", "0.158"],
            ["Decision Tree", "0.516", "0.469", "0.165"],
            ["Naive Bayes", "0.328", "0.316", "0.253"],
            ["Majority baseline", "0.500", "0.333", "n/a"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "Random Forest is best on all three metrics. SVM has the second-best RPS (0.139), which is "
                "interesting: its argmax predictions are not great but its probability outputs are reasonably "
                "calibrated.",
                "para",
            ),
            ("[figures/03_model_comparison.png]", "para"),
            ("Confusion matrix (Random Forest, single seed).", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["", "Pred away", "Pred draw", "Pred home"],
        [
            ["Actual away (n=22)", "17", "2", "3"],
            ["Actual draw (n=10)", "3", "4", "3"],
            ["Actual home (n=32)", "9", "6", "17"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "The biggest off-diagonal mass is \"actual home, predicted away\" (9 cases). This is the "
                "class-imbalance-handling at work: with class_weight='balanced', the model is happy to predict "
                "the smaller class when it sees ambiguous evidence, so it sometimes chooses \"away\" on close "
                "matchups where the home team narrowly wins. Draw recall is only 4/10 = 40%, consistent with "
                "the literature.",
                "para",
            ),
            ("[figures/05_confusion_matrix_best.png]", "para"),
            ("Temporal CV vs Stratified CV gap (information leakage measurement).", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["Model", "Temporal CV", "Stratified 5-Fold", "Gap"],
        [
            ["KNN", "0.478", "0.529", "+0.051"],
            ["Decision Tree", "0.513", "0.607", "+0.094"],
            ["SVM (RBF)", "0.449", "0.510", "+0.061"],
            ["Random Forest", "0.545", "0.617", "+0.071"],
            ["Neural Network", "0.411", "0.598", "+0.186"],
            ["XGBoost", "0.514", "0.611", "+0.097"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "Stratified 5-fold CV overestimates accuracy by 4 to 9 percentage points across all models (and "
                "by 18.6 points for the Neural Network, which is also the most variable). We treat this as one "
                "of our headline methodological findings: shuffled cross-validation is not appropriate for "
                "time-ordered sports data and gives a systematically optimistic view of generalization.",
                "para",
            ),
            ("[figures/04_temporal_vs_stratified_cv.png]", "para"),
        ],
    )

    # 5.3 (added subsection for presentation feedback)
    add_italic_prompt(
        doc,
        "5.3. Test-set breakdown by interpretable groupings (presentation feedback). This subsection "
        "responds to a Part 2 presentation comment about the home-vs-away framing being unhelpful at a "
        "neutral venue, and reframes the same Random Forest predictions.",
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "The reviewer's second follow-up was that \"home wins were predicted best\" is not a useful "
                "claim because at a neutral venue there is no genuine home advantage. The home/away label is "
                "just FIFA's team_1 vs team_2 convention. The reviewer asked for results sliced by something "
                "more interpretable (ELO tier, confederation, prior-champion status) so we can answer the "
                "question \"who should I bet on?\".",
                "para",
            ),
            (
                "We addressed this in scripts/test_set_breakdown.py without retraining. The same Random Forest "
                "predictions on the same 64 test matches are sliced four ways and re-labeled.",
                "para",
            ),
            ("Re-labeling.", "heading"),
            (
                "Instead of {home win, away win, draw}, we use {favorite wins, draw, upset}, where favorite is "
                "the team with the higher pre-match ELO. This is a pure relabeling of the same predictions. To "
                "reduce single-seed noise, we use a 10-seed RF ensemble with averaged class probabilities "
                "(ensemble accuracy 0.594, per-seed range 0.594 to 0.625).",
                "para",
            ),
            ("Reframed confusion matrix.", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["", "Pred upset", "Pred draw", "Pred favorite"],
        [
            ["Actual upset (n=16)", "4", "0", "**12**"],
            ["Actual draw (n=10)", "2", "4", "4"],
            ["Actual favorite (n=38)", "0", "8", "**30**"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "Reading this directly: the model gets 30 of 38 favorite-wins right (79% recall) and only 4 of "
                "16 upsets right (25% recall). It is essentially a \"back the higher-ELO team\" predictor with "
                "confidence inflation. This is a much more honest description of the model than \"home wins were "
                "predicted best\" (which is just the same model under a less informative label).",
                "para",
            ),
            ("[figures/13_favorite_vs_upset_cm.png]", "para"),
            ("Slice breakdowns.", "heading"),
        ],
    )
    add_table_in_box(
        doc,
        ["Slice", "Best cell", "Worst cell", "Comment"],
        [
            ["ELO-tier matchup", "High vs Mid (n=14): 0.71", "Mid vs Low (n=9): 0.44", "Clear-but-not-extreme favorites are easy"],
            ["|ELO diff| bucket", "0-50 toss-up (n=16): 0.69", "200+ \"clear favorite\" (n=10): 0.50", "Counterintuitive (see below)"],
            ["Confederation matchup", "Both other (n=7): 0.71", "Top vs other (n=40): 0.55", "Cross-confederation is upset territory"],
            ["Prior-WC-title matchup", "Champion vs non (n=28): 0.61", "Both prior champs (n=3): 0.33", "Prior titles do not help; n is too small for the worst cell"],
        ],
    )
    add_box_with_paragraphs(
        doc,
        [
            ("[figures/12_test_grouped_metrics.png]", "para"),
            ("The counter-intuitive finding.", "heading"),
            (
                "The 200-plus ELO-gap bucket gives only 0.50 accuracy. The naive expectation is that the bigger "
                "the gap, the easier the prediction. But 2022 contained seven matches with ELO gap >= 150 where "
                "the underdog won outright: Argentina-Saudi Arabia (gap 250), Germany-Japan (gap 257), Tunisia "
                "beat France (gap 286), Cameroon beat Brazil (gap 328), Belgium-Morocco, Japan beat Spain, and "
                "Korea beat Portugal. The model rated all seven as favorite-wins with at least 40% confidence. "
                "Pre-match features simply cannot anticipate these outcomes; they are the residual irreducible "
                "randomness of the sport.",
                "para",
            ),
            ("Knockout-stage betting card.", "heading"),
            (
                "For the 16 knockout matches, the 3-class RF went 12 of 16 (75%). The four misses were "
                "Spain-Morocco (round of 16), Brazil-Croatia (quarterfinal, decided on penalties), "
                "Portugal-Morocco (quarterfinal), and France-Argentina (final, decided on penalties). All four "
                "were upsets predicted as favorite-wins with 67 to 80% confidence.",
                "para",
            ),
            ("[figures/14_knockout_betting_card.png]", "para"),
        ],
    )


def section_6_discussion(doc):
    add_heading(doc, "6. Discussion", level=1)

    # 6.1
    add_italic_prompt(
        doc,
        "6.1. Discussion of the Results. Please provide a thorough and insightful analysis of the "
        "results, integrating them with previous/related work.",
    )
    add_box_with_paragraphs(
        doc,
        [
            (
                "Our tuned Random Forest reaches 0.625 single-seed accuracy and 0.578 macro F1 on the 2022 "
                "World Cup. The 10-seed ensemble accuracy is 0.606, so our most defensible point estimate sits "
                "at 0.61, with a single-seed range of 0.59 to 0.63. The temporal CV F1 of 0.541 over 256 "
                "validation matches across four tournament-year folds gives the more stable performance "
                "estimate. Both numbers are competitive with the 53% to 55% three-class accuracy reported by "
                "Hvattum and Arntzen (2010) on a much larger international-football dataset, although the "
                "comparison is rough because their data and tournament composition differ from ours.",
                "para",
            ),
            ("Regularization beats feature engineering on small data, by roughly 4 to 1.", "heading"),
            (
                "The single most impactful change in the entire project was raising min_samples_leaf from 1 to "
                "5, which improved CV F1 by +0.043. Adding the best new feature (continent advantage) improved "
                "it by +0.010. Adding the second-best feature (qualifying win rate) was net neutral. Several "
                "other external sources (FIFA points, squad market value, international ELO) actively hurt "
                "performance. Mechanistically, the model on 900 rows is variance-limited, not "
                "information-limited. New features either add redundant signal (so the model overfits to which "
                "version of the same signal it sees) or they make the model more confident on the favorite at "
                "the expense of draw recall. Either way, the marginal feature is worth less than the marginal "
                "regularization. This generalizes: on small datasets, tune capacity before chasing features.",
                "para",
            ),
            ("The \"overconfidence\" pattern.", "heading"),
            (
                "Three of the five external feature sources we tested (FIFA rankings, squad market value, "
                "international ELO) shared a failure mode. Each one is a more precise team-strength measure "
                "than World-Cup-only ELO, which only updates every four years. With these added, the model "
                "became confident in its top-pick winner and stopped predicting draws almost entirely. Macro F1 "
                "collapsed because draw F1 went near zero. The mechanism connects to Dixon and Coles (1997): "
                "draws are not a separate process; they are a regime where the two teams' strengths are close "
                "enough that the score-difference distribution straddles zero. A more precise strength measure "
                "narrows that straddle and makes the model bet on a winner. Our blunt World-Cup-only ELO has "
                "the opposite effect, which we describe as \"productive imprecision\". The four-year update gap "
                "means many teams enter a tournament with similar ratings, and the model treats draws as "
                "plausible.",
                "para",
            ),
            ("Domain mismatch: more data did not help.", "heading"),
            (
                "We tried expanding the training set from 900 World Cup matches to 28,000 international matches "
                "(qualifiers, friendlies, continental cups). Every configuration performed worse than the "
                "World-Cup-only baseline. The reason is structural. In qualifiers and friendlies, the home team "
                "plays at their actual home stadium and wins about 60% of the time. In World Cup matches, the "
                "home team is a FIFA administrative label at a neutral venue and wins about 50% of the time. "
                "The model trained on the larger dataset learned a strong home-win prior that does not transfer. "
                "This is a textbook out-of-distribution training problem: more data only helps when the "
                "additional data is from the same distribution as the test data.",
                "para",
            ),
            ("Temporal CV is essential.", "heading"),
            (
                "Across all six models we evaluated under both temporal walk-forward CV and stratified 5-fold "
                "CV, stratified k-fold overestimated accuracy by 4 to 9 percentage points, and 18 points for "
                "the Neural Network. The mechanism is that shuffled k-fold lets the model train on future "
                "tournaments and validate on past ones, which is especially leaky for time-evolving features "
                "like rolling form and expanding-window aggregates. We treat this as one of our headline "
                "methodological contributions and recommend temporal walk-forward CV as the default for any "
                "sports or time-series prediction project.",
                "para",
            ),
            ("The reviewer was right about draws (Section 4.4).", "heading"),
            (
                "Removing draws and predicting binary win-or-loss raised the same Random Forest from 0.625 to "
                "0.704 accuracy and from 0.578 to 0.724 macro F1 on the same 39 features and the same temporal "
                "split. This is a clean +0.146 F1 gain attributable specifically to the binary framing. The "
                "knockout-only variant pushes AUC to 0.821 but the test set is too small (n=16) for that to be "
                "a robust headline. The all-stages binary result is the more credible takeaway: most of the "
                "3-class residual error is concentrated in draws.",
                "para",
            ),
            ("The reframing makes the model's behavior interpretable (Section 5.3).", "heading"),
            (
                "Re-labeling outcomes as favorite, draw, or upset (with favorite = higher pre-match ELO) reveals "
                "that the same RF is fundamentally a \"back the favorite\" predictor. It catches 79% of actual "
                "favorite-wins but only 25% of upsets. The 2022 tournament happened to be an unusually "
                "upset-heavy edition. Seven matches had ELO gaps of at least 150 with the underdog winning "
                "outright, and the model rated all seven as favorite-wins. The \"200+ ELO gap\" bucket gave "
                "only 50% accuracy, contrary to the naive expectation that bigger gaps should be easier. We "
                "think this reframing is more honest than \"home wins are predicted best,\" and we have updated "
                "the relevant slides accordingly.",
                "para",
            ),
            ("Unsupervised analysis as a difficulty check.", "heading"),
            (
                "PCA on the 39 features shows that 15 components explain 90% of variance, so the features are "
                "not redundant. K-Means with k=3 against the actual outcome labels gives an Adjusted Rand Index "
                "of 0.022, which is essentially zero. t-SNE shows local clusters but no separation by outcome "
                "class. In other words, even with perfect knowledge of every feature in advance, the three "
                "outcomes do not form distinct regions of feature space. This sets a hard ceiling on what any "
                "classifier could achieve and reinforces that the residual error is not a model-architecture "
                "problem; it is the irreducible randomness of football.",
                "para",
            ),
            ("[figures/07_pca_explained_variance.png]", "para"),
            ("[figures/09_tsne_scatter.png]", "para"),
        ],
    )

    # 6.2
    add_italic_prompt(
        doc,
        "6.2. Recommendations and implications of your report. Provide relevant recommendations, "
        "develop thoughtful implications, and suggest specific actionable directions for the future.",
    )
    add_box_with_paragraphs(
        doc,
        [
            ("We close with five concrete recommendations and a broader implication.", "para"),
            ("1. Incorporate pre-match betting odds.", "heading"),
            (
                "Closing-line odds aggregate large amounts of information about injuries, lineups, and tactical "
                "matchups that we cannot easily encode as features. Groll et al. (2019) and many others have "
                "shown that odds-based features are competitive with or superior to most feature combinations. "
                "Adding them is the cheapest available accuracy improvement and would let us probe how much of "
                "the residual error after our current features is \"really\" irreducible randomness versus "
                "information we just did not encode.",
                "para",
            ),
            ("2. Try ordinal regression for the three-class problem.", "heading"),
            (
                "Football outcomes have natural ordering (away win < draw < home win). Models that respect this "
                "structure (Hubacek, Sourek, and Zelezny, 2019) might improve draw prediction by treating it as "
                "an intermediate outcome. Right now our multiclass classifiers treat the three classes as "
                "unordered, which throws information away. The implementation is straightforward: replace the "
                "categorical loss with a cumulative-link or score-based ordinal loss.",
                "para",
            ),
            ("3. Prioritize regularization over feature engineering on small datasets.", "heading"),
            (
                "Our finding that a single hyperparameter change (min_samples_leaf 1 to 5) outperformed every "
                "external feature source by 4-to-1 has practical implications beyond this project. Before "
                "investing engineering effort in collecting additional feature sources, tune model capacity. "
                "For tree ensembles, that means min_samples_leaf, max_depth, max_features, and (if using "
                "gradient boosting) learning_rate and reg_lambda. We estimate the \"feature budget\" on 900 "
                "rows is around 37 to 39 features before overfitting starts to dominate any new signal.",
                "para",
            ),
            ("4. Recompute ELO from international matches with a variable K-factor.", "heading"),
            (
                "Our World-Cup-only ELO has the productive-imprecision property that helps draw prediction, but "
                "the 4-year gap between updates is genuinely a limitation. A practical compromise would be to "
                "compute ELO from all international matches (qualifiers, friendlies, continental cups) but with "
                "a much higher K-factor for World Cup matches than for friendlies, so that the system reflects "
                "both the volume of international play and the heightened stakes of tournament matches. This "
                "would have to be paired with careful handling of the home-team domain mismatch we identified.",
                "para",
            ),
            ("5. Revisit StatsBomb expected goals when coverage expands.", "heading"),
            (
                "Among the five external feature sources we tested, expected goals had the lowest correlation "
                "with ELO (r=0.35) and captures genuinely novel information about match quality rather than "
                "team strength. It failed for us solely because of coverage (only 10% of training data has real "
                "values). If event-level data becomes available for additional World Cups beyond 2018 and 2022, "
                "xG should be reconsidered. We expect it would help draw prediction specifically, since xG "
                "distinguishes \"deserved\" results from lucky ones in a way that ELO does not.",
                "para",
            ),
            ("Broader implication.", "heading"),
            (
                "This project is, to some extent, an instructive case study in how to handle a small "
                "time-ordered classification problem with imbalanced classes and a partly misleading label. The "
                "lessons that we think generalize beyond football are: (a) evaluate with a temporal split that "
                "respects the deployment setting, (b) tune capacity before adding features, (c) be skeptical of "
                "\"more data\" if the additional data is from a different distribution, and (d) re-frame "
                "outcomes in domain-meaningful terms rather than dataset-administrative terms. The fourth "
                "lesson is one we learned from the presentation feedback rather than from the literature, and "
                "we suspect it generalizes widely to applied classification problems where the labels carry "
                "baked-in conventions that are not obvious to outside readers.",
                "para",
            ),
        ],
    )


def section_7_code(doc):
    add_heading(doc, "7. Code and Data Availability", level=1)
    add_italic_prompt(doc, "Please provide the link to your GitHub repository here:")
    add_box_with_paragraphs(
        doc,
        [
            ("https://github.com/drdiguglielmo-personal/datascienceproject", "para"),
            ("Reproduction recipe:", "heading"),
            ("pip3 install -r requirements.txt", "code"),
            ("python3 scripts/clean_worldcup.py", "code"),
            ("python3 scripts/feature_engineering.py", "code"),
            ("python3 scripts/binary_no_draw_model.py        # presentation feedback (B)", "code"),
            ("python3 scripts/test_set_breakdown.py          # presentation feedback (C)", "code"),
            ("# Then execute Part2_Models_and_Results.ipynb top to bottom", "code"),
            (
                "All raw and cleaned data files are in files_needed/ and data_clean/. All figures are "
                "regenerated by the scripts and notebook. The repository contains a README.md with a "
                "feature-by-feature description and an experiment log.",
                "para",
            ),
        ],
    )


def references_section(doc):
    add_heading(doc, "References", level=1)
    refs = [
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.",
        "Constantinou, A. C., & Fenton, N. E. (2012). Solving the problem of inadequate scoring rules for assessing probabilistic football forecast models. Journal of Quantitative Analysis in Sports, 8(1).",
        "Dixon, M. J., & Coles, S. G. (1997). Modelling association football scores and inefficiencies in the football betting market. Journal of the Royal Statistical Society: Series C (Applied Statistics), 46(2), 265-280.",
        "Elo, A. E. (1978). The Rating of Chess Players, Past and Present. Arco Publishing.",
        "Fjelstul, J. C. (2021). The Fjelstul World Cup Database v1.2.0. https://github.com/jfjelstul/worldcup",
        "Groll, A., Ley, C., Schauberger, G., & Lock, H. (2019). A hybrid random forest to predict soccer matches in international tournaments. Journal of Quantitative Analysis in Sports, 15(4), 271-287.",
        "Hubacek, O., Sourek, G., & Zelezny, F. (2019). Exploiting sports-betting market using machine learning. International Journal of Forecasting, 35(2), 783-796.",
        "Hvattum, L. M., & Arntzen, H. (2010). Using ELO ratings for match result prediction in association football. International Journal of Forecasting, 26(3), 460-470.",
        "Maher, M. J. (1982). Modelling association football scores. Statistica Neerlandica, 36(3), 109-118.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., et al. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "Tax, N., & Joustra, Y. (2015). Predicting the Dutch football competition using public data: A machine learning approach.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before_pt=2, after_pt=2)
        run = p.add_run(r)
        set_run_font(run)


def members_section(doc):
    add_heading(doc, "Members' Contributions", level=1)
    add_simple_paragraph(
        doc,
        "We adopt the CRediT Taxonomy (https://credit.niso.org) to describe each team member's "
        "individual contributions. All team members reviewed and agreed to the contributions below before "
        "submission.",
    )
    add_spacer(doc, pt=4)

    rows = [
        ["Conceptualization", "Ideas; formulation or evolution of overarching research goals and aims.", "Drew DiGuglielmo, Dan Yoo"],
        ["Data Collection", "Activities to search for, obtain, download datasets.", "Drew DiGuglielmo, Charlie Devine"],
        ["Data Curation", "Management activities to annotate, scrub data and maintain research data for initial use and later reuse.", "Drew DiGuglielmo, Dan Yoo"],
        ["Formal Analysis", "Application of statistical, mathematical, computational, or other formal techniques to analyze or synthesize study data.", "Drew DiGuglielmo, Dan Yoo, Michael Sorenson"],
        ["Investigation", "Conducting research and investigation, specifically performing the experiments.", "Drew DiGuglielmo, Michael Sorenson"],
        ["Methodology", "Development or design of methodology; creation of models.", "Drew DiGuglielmo, Dan Yoo"],
        ["Project Administration", "Management and coordination responsibility for the research activity planning and execution.", "Drew DiGuglielmo"],
        ["Software", "Programming, software development; implementation of computer code and supporting algorithms.", "Drew DiGuglielmo, Dan Yoo"],
        ["Validation", "Verification of overall replication/reproducibility of results.", "Drew DiGuglielmo, Charlie Devine"],
        ["Visualization", "Preparation, creation and presentation of the published work, specifically visualization/data presentation.", "Drew DiGuglielmo, Michael Sorenson"],
        ["Writing, Original Draft Preparation", "Creation and presentation of the published work, specifically writing the initial draft.", "Drew DiGuglielmo"],
        ["Writing, Review and Editing", "Critical review, commentary or revision of the published work.", "Drew DiGuglielmo, Dan Yoo, Michael Sorenson, Charlie Devine"],
    ]
    add_table_in_box(
        doc,
        ["Contributor Role", "Role Definition", "Team Members"],
        rows,
    )
    add_simple_paragraph(
        doc,
        "(Team to verify and edit this contribution table to reflect actual roles before submission.)",
        italic=True,
        pt=10,
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    doc = Document()
    set_default_style(doc)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    build_header(doc)
    section_1_introduction(doc)
    section_2_related_work(doc)
    section_3_data(doc)
    section_4_modeling(doc)
    section_5_evaluation(doc)
    section_6_discussion(doc)
    section_7_code(doc)
    references_section(doc)
    members_section(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
